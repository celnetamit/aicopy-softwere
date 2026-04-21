"""Document processing, AI editing, and DOCX generation."""

import os
import re
import json
import difflib
import html
import subprocess
import tempfile
import zipfile
import requests
from typing import Tuple, List, Dict, Optional, Set
from collections import Counter
from xml.etree import ElementTree as ET
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from chicago_editor import ChicagoEditor


class DocumentProcessor:
    """Handles document loading, AI processing, and DOCX output."""

    MISSING_PLACEHOLDER_COLOR = RGBColor(128, 128, 128)
    CELL_PARAGRAPH_MARKER = "[[CELL_PARA]]"
    CELL_BLOCK_MARKER = "[[CELL_BLOCK]]"
    CELL_TABLE_ROW_MARKER = "[[CELL_TABLE_ROW]]"
    CELL_TABLE_CELL_MARKER = "[[CELL_TABLE_CELL]]"
    TEXTBOX_MARKER = "[[TEXTBOX]]"
    TEXTBOX_PARAGRAPH_MARKER = "[[TEXTBOX_PARA]]"
    HEADING_LINE_RE = re.compile(r'^\s{0,3}(#{1,6})\s+(.*\S)\s*$')
    BULLET_LINE_RE = re.compile(r'^\s*([-*•])\s+(.*\S)\s*$')
    NUMBERED_LINE_RE = re.compile(r'^\s*((?:\d+|[A-Za-z])[.)])\s+(.*\S)\s*$')
    PLAINTEXT_HEADING_RE = re.compile(r'^[A-Z][A-Za-z0-9/&,\- ]{0,79}:?$')

    def __init__(self, ollama_host: str = "http://localhost:11434"):
        self.ollama_host = ollama_host
        self.editor = ChicagoEditor()
        self.model = "llama3.1"
        self.gemini_model = "gemini-1.5-flash"
        self.openrouter_model = "openrouter/auto"
        self.agent_router_model = "gpt-5"
        self._last_ai_warning = ""
        self._last_selection_note = ""
        self._last_ai_pipeline_note = ""
        self._last_chunk_decisions: List[Dict] = []
        self._last_processing_audit: Dict = {}
        self._last_journal_profile_report: Dict = {}
        self._last_citation_reference_report: Dict = {}
        self._last_docx_package_report: Dict = {}
        self._nlp = None

    def _reset_processing_audit(self):
        """Reset per-run audit data."""
        self._last_chunk_decisions = []
        self._last_processing_audit = {
            "mode": "rule_only",
            "sections": [],
            "summary": {},
        }
        self._last_journal_profile_report = {}
        self._last_citation_reference_report = {}

    def _quality_score(self, risk_score: Optional[int]) -> int:
        """Convert risk score (lower is better) into quality score (0-100)."""
        if risk_score is None:
            return 0
        return max(0, 100 - min(100, int(risk_score)))

    def _risk_label(self, risk_score: Optional[int]) -> str:
        """Return label for risk score."""
        if risk_score is None:
            return "unknown"
        score = int(risk_score)
        if score <= 8:
            return "low"
        if score <= 25:
            return "medium"
        return "high"

    def load_document(self, path: str) -> Tuple[str, str]:
        """Load document and return (text, format)."""
        ext = os.path.splitext(path)[1].lower()

        if ext == '.txt':
            self._last_docx_package_report = {}
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            return text, 'txt'

        elif ext == '.docx':
            doc = Document(path)
            text_parts = [block["text"] for block in self._extract_docx_blocks(doc) if block.get("consumes_text", True)]
            self._last_docx_package_report = self._inspect_docx_package(path)
            self._attach_docx_package_summary()
            text = '\n'.join(text_parts)
            return text, 'docx'

        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def process_text(self, text: str, options: Dict) -> str:
        """Process text using AI with Chicago style rules."""
        self._last_selection_note = ""
        self._last_ai_pipeline_note = ""
        self._reset_processing_audit()
        self._attach_docx_package_summary()

        # First apply rule-based corrections
        rules_corrected = self.editor.correct_all(text, options)

        # Then enhance with AI for context-aware corrections
        ai_corrected = self._call_ai_editor(text, rules_corrected, options)
        if ai_corrected:
            ai_first_cmos = self._is_ai_first_cmos_mode(options)
            # In AI-first CMOS mode, keep AI as the language authority and only
            # apply structure-safe normalization afterward.
            ai_post_processed = self._postprocess_ai_output(ai_corrected, options)
            selected = self._select_best_correction(
                original=text,
                rules_corrected=rules_corrected,
                ai_corrected=ai_post_processed,
                prefer_ai=ai_first_cmos,
            )
            if self._last_ai_pipeline_note:
                self._last_selection_note = f"{self._last_ai_pipeline_note} {self._last_selection_note}".strip()
            self._attach_cmos_guardrails(original=text, corrected=selected, options=options)
            self._last_journal_profile_report = self.editor.build_reference_profile_report(selected, options)
            self._last_citation_reference_report = self.editor.build_citation_reference_validator_report(selected, options)
            return selected

        if self._last_ai_pipeline_note:
            self._last_selection_note = self._last_ai_pipeline_note
        else:
            self._last_selection_note = "Rule-based correction applied."
        if self._last_processing_audit.get("mode") == "full":
            self._last_processing_audit["summary"] = {"reason": "AI full-pass returned no usable output"}
        else:
            self._last_processing_audit["mode"] = "rule_only"
            self._last_processing_audit["summary"] = {"reason": "AI disabled or unavailable"}
        self._attach_cmos_guardrails(original=text, corrected=rules_corrected, options=options)
        self._last_journal_profile_report = self.editor.build_reference_profile_report(rules_corrected, options)
        self._last_citation_reference_report = self.editor.build_citation_reference_validator_report(rules_corrected, options)
        return rules_corrected

    def _is_ai_first_cmos_mode(self, options: Dict) -> bool:
        ai_options = options.get("ai", {}) if isinstance(options, dict) else {}
        if not isinstance(ai_options, dict):
            return False
        return bool(ai_options.get("ai_first_cmos", False))

    def _postprocess_ai_output(self, ai_text: str, options: Dict) -> str:
        if self._is_ai_first_cmos_mode(options):
            return self.editor.postprocess_ai_first_copyedit(ai_text, options)
        return self.editor.correct_all(ai_text, options)

    def _attach_cmos_guardrails(self, original: str, corrected: str, options: Dict):
        """Attach CMOS workflow guardrails summary to processing audit."""
        safe_options = options if isinstance(options, dict) else {}
        domain_report = self.get_domain_report()
        detected_domain = str(domain_report.get("profile", "general") or "general").strip().lower()
        requested_domain = str(safe_options.get("domain_profile", "auto") or "auto").strip().lower()
        protected_terms = int(domain_report.get("protected_terms", 0) or 0)
        custom_terms = int(domain_report.get("custom_terms", 0) or 0)
        chicago_enabled = bool(safe_options.get("chicago_style", True))
        cmos_strict_mode = bool(safe_options.get("cmos_strict_mode", True))

        ai_options = safe_options.get("ai", {}) if isinstance(safe_options.get("ai"), dict) else {}
        ai_enabled = bool(ai_options.get("enabled", False))
        ai_provider = str(ai_options.get("provider", "ollama") or "ollama").strip().lower()

        warnings: List[str] = []
        recommendations: List[str] = []

        if not chicago_enabled:
            warnings.append("Chicago style option is disabled, so CMOS compliance is limited.")

        if requested_domain == "auto":
            recommendations.append(
                "Set an explicit domain profile (medical/law/engineering/general) for stricter term protection."
            )
            if cmos_strict_mode and detected_domain in ("medical", "law", "engineering"):
                warnings.append(
                    f"Detected domain is {detected_domain}; auto profile may be too loose for strict CMOS editing."
                )

        if cmos_strict_mode and detected_domain in ("medical", "law", "engineering"):
            if custom_terms < 5:
                warnings.append(
                    f"Only {custom_terms} custom glossary terms configured for {detected_domain} text."
                )
                recommendations.append("Add key domain terms to Custom Terms to avoid unsafe rewrites.")
            if protected_terms < 25:
                warnings.append(
                    f"Protected term coverage is low ({protected_terms}); increase dictionary/glossary coverage."
                )

        if cmos_strict_mode and len(str(original or "")) > 12000 and not ai_enabled:
            recommendations.append("Enable AI section-wise mode for long manuscripts, then review audit fallback reasons.")

        if ai_enabled and ai_provider == "openrouter":
            recommendations.append("Validate OpenRouter key/model in Admin Dashboard before production runs.")
        if ai_enabled and ai_provider == "agent_router":
            recommendations.append("Validate AgentRouter token/model in Admin Dashboard before production runs.")

        compliance_score = 100
        compliance_score -= min(60, len(warnings) * 12)
        compliance_score -= 8 if not chicago_enabled else 0
        compliance_score -= 8 if requested_domain == "auto" and cmos_strict_mode else 0
        compliance_score = max(0, min(100, compliance_score))

        if compliance_score >= 85:
            status = "strong"
        elif compliance_score >= 65:
            status = "needs_attention"
        else:
            status = "at_risk"

        guardrails = {
            "strict_mode": cmos_strict_mode,
            "status": status,
            "compliance_score": compliance_score,
            "requested_domain": requested_domain,
            "detected_domain": detected_domain,
            "protected_terms": protected_terms,
            "custom_terms": custom_terms,
            "chicago_style_enabled": chicago_enabled,
            "ai_enabled": ai_enabled,
            "warnings": warnings,
            "recommendations": recommendations,
            "input_chars": len(str(original or "")),
            "output_chars": len(str(corrected or "")),
        }

        summary = self._last_processing_audit.setdefault("summary", {})
        if not isinstance(summary, dict):
            summary = {}
            self._last_processing_audit["summary"] = summary
        summary["cmos_guardrails"] = guardrails

    def _attach_docx_package_summary(self):
        """Attach DOCX package preservation details to the processing audit summary."""
        if not self._last_docx_package_report:
            return
        summary = self._last_processing_audit.setdefault("summary", {})
        if not isinstance(summary, dict):
            summary = {}
            self._last_processing_audit["summary"] = summary
        summary["docx_package_features"] = dict(self._last_docx_package_report)

        present = []
        for key in ("comments", "footnotes", "endnotes", "textboxes"):
            count = int(self._last_docx_package_report.get(key, 0) or 0)
            if count > 0:
                present.append(f"{key}={count}")
        if present:
            note = "DOCX package contains " + ", ".join(present) + "; export preserves these structures, but editing remains body-text-first."
            if note not in self._last_selection_note:
                self._last_selection_note = f"{self._last_selection_note} {note}".strip()

    def _call_ai_editor(self, original: str, rules_corrected: str, options: Dict) -> Optional[str]:
        """Call the configured AI provider for enhanced editing."""
        settings = self._get_ai_settings(options)
        if not settings["enabled"]:
            self._last_processing_audit["mode"] = "rule_only"
            self._last_processing_audit["summary"] = {"reason": "AI disabled in settings"}
            return None

        if self._should_use_section_analysis(original, rules_corrected, options):
            return self._call_ai_editor_sectioned(original, rules_corrected, options, settings)

        self._last_processing_audit["mode"] = "full"
        prompt = self._build_edit_prompt(original, rules_corrected, options, stage="full")
        return self._invoke_ai_provider(prompt, settings)

    def _invoke_ai_provider(self, prompt: str, settings: Dict) -> Optional[str]:
        """Dispatch AI request to configured provider."""
        if settings["provider"] == "gemini":
            return self._call_gemini_editor(prompt, settings)
        if settings["provider"] == "openrouter":
            return self._call_openrouter_editor(prompt, settings)
        if settings["provider"] == "agent_router":
            return self._call_agent_router_editor(prompt, settings)
        return self._call_ollama_editor(prompt, settings)

    def _should_use_section_analysis(self, original: str, rules_corrected: str, options: Dict) -> bool:
        """Decide whether manuscript should be processed section-wise."""
        ai_options = options.get("ai", {}) if isinstance(options, dict) else {}
        if not isinstance(ai_options, dict):
            ai_options = {}

        if ai_options.get("section_wise") is False:
            return False

        source = rules_corrected or original or ""
        threshold_chars = ai_options.get("section_threshold_chars", 12000)
        threshold_paragraphs = ai_options.get("section_threshold_paragraphs", 90)
        try:
            threshold_chars = max(4000, int(threshold_chars))
        except Exception:
            threshold_chars = 12000
        try:
            threshold_paragraphs = max(20, int(threshold_paragraphs))
        except Exception:
            threshold_paragraphs = 90

        paragraph_count = source.count('\n') + 1
        return len(source) > threshold_chars or paragraph_count > threshold_paragraphs

    def _is_major_heading_line(self, line: str) -> bool:
        """Heuristic heading detection used to keep chunk boundaries semantically clean."""
        value = (line or "").strip()
        if not value:
            return False
        if re.fullmatch(
            r'(?i)(abstract|introduction|background|method(?:s|ology)?|results?|discussion|conclusion|references?)',
            value
        ):
            return True
        if len(value) < 110 and value == value.upper() and re.search(r'[A-Z]', value):
            return True
        if len(value) < 95 and re.fullmatch(r'[A-Z][A-Za-z0-9,&:()\'’\-./ ]*', value) and not re.search(r'[.!?]$', value):
            return True
        return False

    def _split_for_section_analysis(
        self,
        original: str,
        rules_corrected: str,
        options: Dict
    ) -> List[Dict]:
        """Split manuscript into aligned sections for long-context processing."""
        ai_options = options.get("ai", {}) if isinstance(options, dict) else {}
        if not isinstance(ai_options, dict):
            ai_options = {}

        max_chars = ai_options.get("section_chunk_chars", 5500)
        max_lines = ai_options.get("section_chunk_lines", 28)
        try:
            max_chars = max(1800, int(max_chars))
        except Exception:
            max_chars = 5500
        try:
            max_lines = max(8, int(max_lines))
        except Exception:
            max_lines = 28

        orig_lines = (original or "").split('\n')
        base_lines = (rules_corrected or "").split('\n')
        line_count = max(len(orig_lines), len(base_lines))

        if line_count == 0:
            return []

        heading_re = re.compile(r'^\s*references?\s*:?\s*$', flags=re.IGNORECASE)
        section_break_re = re.compile(
            r'^\s*(?:appendix|acknowledg(?:e)?ments?|funding|conflicts?\s+of\s+interest|author\s+contributions?)\s*:?\s*$',
            flags=re.IGNORECASE
        )

        chunks: List[Dict] = []
        bucket: List[Tuple[str, str]] = []
        bucket_start = 0
        bucket_chars = 0
        in_references = False

        def flush(end_index: int):
            nonlocal bucket, bucket_start, bucket_chars
            if not bucket:
                return
            chunks.append({
                "index": len(chunks) + 1,
                "line_start": bucket_start + 1,
                "line_end": end_index + 1,
                "original": '\n'.join(item[0] for item in bucket),
                "baseline": '\n'.join(item[1] for item in bucket),
            })
            bucket = []
            bucket_chars = 0

        for i in range(line_count):
            orig_line = orig_lines[i] if i < len(orig_lines) else ""
            base_line = base_lines[i] if i < len(base_lines) else ""
            check_line = (base_line or orig_line).strip()

            if heading_re.match(check_line):
                in_references = True
            elif in_references and section_break_re.match(check_line):
                in_references = False

            line_len = max(len(orig_line), len(base_line)) + 1
            should_split = False
            if bucket:
                too_big = (bucket_chars + line_len > max_chars) or (len(bucket) >= max_lines)
                heading_boundary = self._is_major_heading_line(check_line) and len(bucket) >= 6
                if not in_references and (too_big or heading_boundary):
                    should_split = True

            if should_split:
                flush(i - 1)
                bucket_start = i

            if not bucket:
                bucket_start = i

            bucket.append((orig_line, base_line))
            bucket_chars += line_len

        flush(line_count - 1)
        return chunks

    def _choose_candidate(
        self,
        original: str,
        baseline: str,
        ai_candidate: str,
        tolerance: int = 3,
        stage: str = "section",
        prefer_ai: bool = False,
    ) -> Tuple[str, bool, Dict]:
        """Choose safer text between baseline and AI candidate using risk scoring."""
        base_score, base_reasons = self._candidate_risk_score(original, baseline)
        detail: Dict = {
            "stage": stage,
            "tolerance": max(0, int(tolerance)),
            "prefer_ai": bool(prefer_ai),
            "baseline_score": base_score,
            "baseline_reasons": base_reasons,
            "baseline_quality": self._quality_score(base_score),
            "baseline_risk": self._risk_label(base_score),
            "ai_score": None,
            "ai_reasons": [],
            "ai_quality": 0,
            "ai_risk": "unknown",
            "score_delta": None,
            "score_margin": None,
            "decision": "fallback",
            "decision_reason": "empty ai output",
            "decision_confidence": "low",
            "accepted_ai": False,
        }

        if not (ai_candidate or "").strip():
            return baseline, False, detail

        ai_score, ai_reasons = self._candidate_risk_score(original, ai_candidate)
        detail["ai_score"] = ai_score
        detail["ai_reasons"] = ai_reasons
        detail["ai_quality"] = self._quality_score(ai_score)
        detail["ai_risk"] = self._risk_label(ai_score)
        detail["score_delta"] = ai_score - base_score
        detail["score_margin"] = (base_score + detail["tolerance"]) - ai_score

        accept_tolerance = max(0, int(tolerance))
        hard_risk_limit = 12 if prefer_ai else None
        if hard_risk_limit is not None:
            accept_tolerance = max(accept_tolerance, 6)

        if ai_score <= base_score + accept_tolerance and (hard_risk_limit is None or ai_score <= hard_risk_limit):
            detail["decision"] = "accepted"
            if hard_risk_limit is not None:
                detail["decision_reason"] = (
                    f"ai-first mode accepted ai score {ai_score} with baseline {base_score}, "
                    f"tolerance {accept_tolerance}, hard limit {hard_risk_limit}"
                )
            else:
                detail["decision_reason"] = f"ai score {ai_score} <= baseline {base_score} + tolerance {detail['tolerance']}"
            detail["decision_confidence"] = "high" if ai_score <= base_score else "medium"
            detail["accepted_ai"] = True
            return ai_candidate, True, detail

        detail["decision"] = "fallback"
        if hard_risk_limit is not None and ai_score > hard_risk_limit:
            detail["decision_reason"] = f"ai-first mode rejected ai score {ai_score} > hard limit {hard_risk_limit}"
        else:
            detail["decision_reason"] = f"ai score {ai_score} > baseline {base_score} + tolerance {detail['tolerance']}"
        detail["decision_confidence"] = "high" if (detail["score_margin"] or 0) < -8 else "medium"
        return baseline, False, detail

    def _build_consistency_prompt(self, manuscript_text: str, options: Dict) -> str:
        """Build a lightweight global consistency prompt over already corrected text."""
        requested_domain = str(options.get('domain_profile', 'auto')).strip().lower()
        detected_domain = self.get_domain_report().get("profile", "general")
        domain_for_prompt = detected_domain if requested_domain == "auto" else requested_domain
        if domain_for_prompt not in ("general", "medical", "engineering", "law"):
            domain_for_prompt = "general"

        journal_profile = self.editor.resolve_journal_profile(options)
        ai_first_cmos = self._is_ai_first_cmos_mode(options)
        initials_rule = "without periods (Smith AB)" if not bool(journal_profile.get("initials_with_periods")) else "with periods (Smith A.B.)"
        title_rule = "sentence case" if str(journal_profile.get("title_case", "sentence")) == "sentence" else "title case"
        journal_rule = "NLM abbreviations" if str(journal_profile.get("journal_abbrev", "nlm")) == "nlm" else "full journal names"

        extra_goal = (
            "- Use professional CMOS judgment for grammar and capitalization; avoid unnecessary rewrites.\n"
            if ai_first_cmos else
            ""
        )

        return f"""You are a senior copy editor doing a final consistency pass.
The manuscript was already corrected section-by-section.
Return ONLY the revised manuscript text.

Goals:
- Keep edits minimal; do not rewrite content.
- Harmonize terminology and style consistently across all sections.
- Preserve paragraph structure and heading order.
- Keep citations numeric in brackets, e.g. [9, 19].
- Ensure citations and references follow Vancouver first-appearance numbering.
- Use journal profile "{journal_profile.get('label', 'Vancouver')}".
- In references, use initials {initials_rule}, titles in {title_rule}, and journal names as {journal_rule}.
- Preserve URLs/DOIs/emails exactly.
- Preserve {domain_for_prompt} technical terms.
{extra_goal}

Manuscript:
{manuscript_text}

Final consistent manuscript:"""

    def _call_ai_editor_sectioned(self, original: str, rules_corrected: str, options: Dict, settings: Dict) -> Optional[str]:
        """Run section-wise AI correction with safety checks and optional global pass."""
        chunks = self._split_for_section_analysis(original, rules_corrected, options)
        if len(chunks) <= 1:
            self._last_processing_audit["mode"] = "full"
            prompt = self._build_edit_prompt(original, rules_corrected, options, stage="full")
            return self._invoke_ai_provider(prompt, settings)

        ai_options = options.get("ai", {}) if isinstance(options, dict) else {}
        if not isinstance(ai_options, dict):
            ai_options = {}

        def clamp_int(raw_value, default_value, min_value, max_value):
            try:
                parsed = int(raw_value)
            except Exception:
                parsed = int(default_value)
            return max(int(min_value), min(int(max_value), parsed))

        def avg(values: List[int]) -> Optional[float]:
            if not values:
                return None
            return round(sum(values) / len(values), 2)

        def normalize_reason(reason: str) -> str:
            value = (reason or "").strip().lower()
            value = re.sub(r'[^a-z0-9]+', '_', value).strip('_')
            return value or "unknown"

        section_tolerance = clamp_int(ai_options.get("section_accept_tolerance", 4), 4, 0, 12)
        consistency_tolerance = clamp_int(ai_options.get("consistency_tolerance", 3), 3, 0, 12)
        max_consistency_chars = clamp_int(
            ai_options.get("global_consistency_max_chars", 18000),
            18000,
            6000,
            120000,
        )

        outputs: List[str] = []
        accepted_chunks = 0
        total_chunks = len(chunks)
        decisions: List[Dict] = []
        fallback_reason_counts: Dict[str, int] = {}
        self._last_processing_audit["mode"] = "sectioned"

        print(
            "[AI] sectioned analysis start: "
            f"sections={total_chunks} section_tolerance={section_tolerance} "
            f"consistency_tolerance={consistency_tolerance}"
        )

        for idx, chunk in enumerate(chunks, start=1):
            prompt = self._build_edit_prompt(
                chunk["original"],
                chunk["baseline"],
                options,
                stage="section",
                section_index=idx,
                section_total=total_chunks
            )
            ai_chunk = self._invoke_ai_provider(prompt, settings)
            if ai_chunk:
                ai_chunk = self._postprocess_ai_output(ai_chunk, options)
            selected_chunk, used_ai, decision = self._choose_candidate(
                original=chunk["original"],
                baseline=chunk["baseline"],
                ai_candidate=ai_chunk or "",
                tolerance=section_tolerance + (2 if settings.get("ai_first_cmos") else 0),
                stage="section",
                prefer_ai=bool(settings.get("ai_first_cmos")),
            )
            decision["section_index"] = idx
            decision["section_total"] = total_chunks
            decision["line_start"] = int(chunk.get("line_start", 0))
            decision["line_end"] = int(chunk.get("line_end", 0))
            decision["baseline_chars"] = len(chunk.get("baseline", ""))
            decision["ai_chars"] = len(ai_chunk or "")
            decision["selected_chars"] = len(selected_chunk or "")
            decisions.append(decision)

            if used_ai:
                accepted_chunks += 1
            else:
                if decision.get("ai_score") is None:
                    reason_key = "empty_ai_output"
                else:
                    ai_reasons = decision.get("ai_reasons", [])
                    reason_key = normalize_reason(ai_reasons[0] if ai_reasons else "higher_risk_than_baseline")
                fallback_reason_counts[reason_key] = int(fallback_reason_counts.get(reason_key, 0)) + 1

            outputs.append(selected_chunk)

            ai_score = decision.get("ai_score")
            ai_score_text = "None" if ai_score is None else str(ai_score)
            reason = decision.get("decision_reason", "")
            print(
                f"[AI][Section {idx}/{total_chunks}] decision={decision.get('decision')} "
                f"baseline={decision.get('baseline_score')} ai={ai_score_text} "
                f"risk={decision.get('ai_risk')} reason={reason}"
            )

        merged = '\n'.join(outputs)
        if not merged.strip():
            self._last_ai_pipeline_note = "Section-wise AI returned no usable text; using rule-based correction."
            self._last_chunk_decisions = decisions
            self._last_processing_audit["sections"] = decisions
            self._last_processing_audit["summary"] = {
                "total_sections": total_chunks,
                "accepted_sections": accepted_chunks,
                "fallback_sections": max(0, total_chunks - accepted_chunks),
                "acceptance_rate": round((accepted_chunks / total_chunks) * 100, 2) if total_chunks else 0.0,
                "section_tolerance": section_tolerance,
                "consistency_tolerance": consistency_tolerance,
                "fallback_reason_counts": dict(sorted(fallback_reason_counts.items(), key=lambda pair: (-pair[1], pair[0]))),
                "reason": "empty merged output",
            }
            return None

        consistency_audit: Dict = {
            "ran": False,
            "decision": "skipped",
            "reason": "Global consistency pass skipped (manuscript too long).",
            "threshold_chars": max_consistency_chars,
            "merged_chars": len(merged),
        }
        consistency_note = "Global consistency pass skipped (manuscript too long)."
        if len(merged) <= max_consistency_chars:
            consistency_prompt = self._build_consistency_prompt(merged, options)
            ai_consistent = self._invoke_ai_provider(consistency_prompt, settings)
            if ai_consistent:
                ai_consistent = self._postprocess_ai_output(ai_consistent, options)
            merged_selected, used_ai, consistency_decision = self._choose_candidate(
                original=original,
                baseline=merged,
                ai_candidate=ai_consistent or "",
                tolerance=consistency_tolerance + (2 if settings.get("ai_first_cmos") else 0),
                stage="global_consistency",
                prefer_ai=bool(settings.get("ai_first_cmos")),
            )
            merged = merged_selected
            consistency_quality_before = consistency_decision.get("baseline_quality")
            consistency_quality_after = (
                consistency_decision.get("ai_quality")
                if consistency_decision.get("ai_score") is not None
                else None
            )
            quality_delta = None
            if isinstance(consistency_quality_before, int) and isinstance(consistency_quality_after, int):
                quality_delta = consistency_quality_after - consistency_quality_before
            consistency_audit = {
                "ran": True,
                "decision": consistency_decision.get("decision"),
                "reason": consistency_decision.get("decision_reason", ""),
                "tolerance": consistency_tolerance,
                "baseline_score": consistency_decision.get("baseline_score"),
                "ai_score": consistency_decision.get("ai_score"),
                "baseline_quality": consistency_quality_before,
                "ai_quality": consistency_quality_after,
                "quality_delta": quality_delta,
                "details": consistency_decision,
            }
            print(
                "[AI][Consistency] "
                f"decision={consistency_decision.get('decision')} "
                f"baseline={consistency_decision.get('baseline_score')} "
                f"ai={consistency_decision.get('ai_score')}"
            )
            consistency_note = (
                "Global consistency pass applied."
                if used_ai else
                "Global consistency pass returned unstable output; kept section-merged text."
            )

        fallback_chunks = max(0, total_chunks - accepted_chunks)
        baseline_scores: List[int] = [int(d.get("baseline_score", 0)) for d in decisions]
        ai_scores: List[int] = [int(d.get("ai_score")) for d in decisions if isinstance(d.get("ai_score"), int)]
        accepted_ai_scores: List[int] = [
            int(d.get("ai_score"))
            for d in decisions
            if d.get("decision") == "accepted" and isinstance(d.get("ai_score"), int)
        ]
        baseline_quality_scores: List[int] = [int(d.get("baseline_quality", 0)) for d in decisions]
        ai_quality_scores: List[int] = [int(d.get("ai_quality", 0)) for d in decisions if isinstance(d.get("ai_score"), int)]
        accepted_ai_quality_scores: List[int] = [
            int(d.get("ai_quality", 0))
            for d in decisions
            if d.get("decision") == "accepted" and isinstance(d.get("ai_score"), int)
        ]

        sorted_fallback_reason_counts = dict(
            sorted(fallback_reason_counts.items(), key=lambda pair: (-pair[1], pair[0]))
        )

        self._last_chunk_decisions = decisions
        self._last_processing_audit["sections"] = decisions
        self._last_processing_audit["summary"] = {
            "total_sections": total_chunks,
            "accepted_sections": accepted_chunks,
            "fallback_sections": fallback_chunks,
            "acceptance_rate": round((accepted_chunks / total_chunks) * 100, 2) if total_chunks else 0.0,
            "section_tolerance": section_tolerance,
            "consistency_tolerance": consistency_tolerance,
            "average_baseline_risk_score": avg(baseline_scores),
            "average_ai_risk_score": avg(ai_scores),
            "average_accepted_ai_risk_score": avg(accepted_ai_scores),
            "average_baseline_quality": avg(baseline_quality_scores),
            "average_ai_quality": avg(ai_quality_scores),
            "average_accepted_ai_quality": avg(accepted_ai_quality_scores),
            "fallback_reason_counts": sorted_fallback_reason_counts,
            "consistency": consistency_audit,
        }

        self._last_ai_pipeline_note = (
            f"Section-wise AI analysis used ({total_chunks} sections, accepted {accepted_chunks}, fallback {fallback_chunks}, "
            f"acceptance {round((accepted_chunks / total_chunks) * 100, 2) if total_chunks else 0.0}%). "
            f"{consistency_note}"
        )
        return merged

    def _get_ai_settings(self, options: Dict) -> Dict:
        """Get AI settings from request options."""
        ai_options = options.get("ai", {}) if isinstance(options, dict) else {}
        if not isinstance(ai_options, dict):
            ai_options = {}

        provider = str(ai_options.get("provider", "ollama")).strip().lower()
        if provider not in ("ollama", "gemini", "openrouter", "agent_router"):
            provider = "ollama"

        model = str(ai_options.get("model", "")).strip()
        gemini_api_key = str(
            ai_options.get("gemini_api_key", ai_options.get("api_key", os.getenv("GEMINI_API_KEY", "")))
        ).strip()
        openrouter_api_key = str(
            ai_options.get("openrouter_api_key", os.getenv("OPENROUTER_API_KEY", ""))
        ).strip()
        agent_router_api_key = str(
            ai_options.get("agent_router_api_key", os.getenv("AGENT_ROUTER_TOKEN", ""))
        ).strip()

        if provider == "gemini":
            default_model = self.gemini_model
        elif provider == "openrouter":
            default_model = self.openrouter_model
        elif provider == "agent_router":
            default_model = self.agent_router_model
        else:
            default_model = self.model

        return {
            "enabled": bool(ai_options.get("enabled", True)),
            "provider": provider,
            "ollama_host": str(ai_options.get("ollama_host", self.ollama_host)).strip() or self.ollama_host,
            "model": model or default_model,
            "gemini_api_key": gemini_api_key,
            "openrouter_api_key": openrouter_api_key,
            "agent_router_api_key": agent_router_api_key,
            "ai_first_cmos": bool(ai_options.get("ai_first_cmos", False)),
        }

    def _call_ollama_editor(self, prompt: str, settings: Dict) -> Optional[str]:
        """Call Ollama AI for enhanced editing."""
        ollama_host = settings["ollama_host"]
        if not self._check_ollama(ollama_host):
            self._warn_once("Ollama is not reachable; falling back to rule-based editing.")
            return None

        resolved_model = self._resolve_ollama_model(ollama_host, settings["model"])
        if not resolved_model:
            self._warn_once("No Ollama model found; falling back to rule-based editing.")
            return None

        if resolved_model != settings["model"]:
            self._warn_once(
                f"Ollama model '{settings['model']}' not available. Using '{resolved_model}' instead."
            )

        try:
            response = requests.post(
                f"{ollama_host}/api/generate",
                json={
                    "model": resolved_model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {"temperature": 0.3}
                },
                timeout=60
            )

            if response.status_code == 200:
                self._last_ai_warning = ""
                result = response.json().get("response", "")
                return self._extract_corrected_text(result)
            if response.status_code == 404 and "not found" in response.text.lower():
                fallback_model = self._resolve_ollama_model(ollama_host, "")
                if fallback_model and fallback_model != resolved_model:
                    retry = requests.post(
                        f"{ollama_host}/api/generate",
                        json={
                            "model": fallback_model,
                            "prompt": prompt,
                            "stream": False,
                            "options": {"temperature": 0.3}
                        },
                        timeout=60
                    )
                    if retry.status_code == 200:
                        self._warn_once(
                            f"Ollama model '{resolved_model}' not found. Retried with '{fallback_model}'."
                        )
                        self._last_ai_warning = ""
                        result = retry.json().get("response", "")
                        return self._extract_corrected_text(result)
                self._warn_once(
                    f"Ollama model '{resolved_model}' not found; falling back to rule-based editing."
                )
                return None

            self._warn_once(
                f"Ollama editing failed ({response.status_code}); falling back to rule-based editing."
            )
        except Exception as e:
            self._warn_once(f"Ollama editing failed: {e}")

        return None

    def _call_gemini_editor(self, prompt: str, settings: Dict) -> Optional[str]:
        """Call Gemini AI for enhanced editing."""
        if not settings["gemini_api_key"]:
            self._warn_once("Gemini API key not set; falling back to rule-based editing.")
            return None

        model = settings["model"]
        url = (
            f"https://generativelanguage.googleapis.com/v1beta/models/"
            f"{model}:generateContent?key={settings['gemini_api_key']}"
        )

        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.3},
        }

        try:
            response = requests.post(url, json=payload, timeout=60)
            if response.status_code == 200:
                self._last_ai_warning = ""
                return self._extract_corrected_text(self._extract_gemini_text(response.json()))
            if response.status_code == 403:
                self._warn_once(
                    "Gemini request blocked (403). Check API key/project restrictions or use Ollama."
                )
            else:
                self._warn_once(
                    f"Gemini editing failed ({response.status_code}); falling back to rule-based editing."
                )
        except Exception as e:
            self._warn_once(f"Gemini editing failed: {e}")

        return None

    def _call_openrouter_editor(self, prompt: str, settings: Dict) -> Optional[str]:
        """Call OpenRouter-compatible chat completion API."""
        if not settings["openrouter_api_key"]:
            self._warn_once("OpenRouter API key not set; falling back to rule-based editing.")
            return None

        model = settings["model"] or self.openrouter_model
        url = "https://openrouter.ai/api/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {settings['openrouter_api_key']}",
            "Content-Type": "application/json",
        }

        # Optional but recommended by OpenRouter for analytics/rate limits.
        referer = os.getenv("OPENROUTER_HTTP_REFERER", "").strip()
        title = os.getenv("OPENROUTER_APP_TITLE", "Manuscript Editor").strip()
        if referer:
            headers["HTTP-Referer"] = referer
        if title:
            headers["X-Title"] = title

        payload = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": "You are a professional copy editor. Return only corrected manuscript text.",
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.3,
        }

        try:
            response = requests.post(url, headers=headers, json=payload, timeout=75)
            if response.status_code == 200:
                self._last_ai_warning = ""
                return self._extract_corrected_text(self._extract_openrouter_text(response.json()))
            if response.status_code in (401, 403):
                self._warn_once("OpenRouter request unauthorized/forbidden; check API key and model access.")
            elif response.status_code == 429:
                self._warn_once("OpenRouter rate-limited (429); falling back to rule-based editing.")
            else:
                self._warn_once(
                    f"OpenRouter editing failed ({response.status_code}); falling back to rule-based editing."
                )
        except Exception as e:
            self._warn_once(f"OpenRouter editing failed: {e}")

        return None

    def _call_agent_router_editor(self, prompt: str, settings: Dict) -> Optional[str]:
        """Call AgentRouter's OpenAI-compatible chat completion API."""
        if not settings["agent_router_api_key"]:
            self._warn_once("AgentRouter token not set; falling back to rule-based editing.")
            return None

        model = settings["model"] or self.agent_router_model
        url = "https://agentrouter.org/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {settings['agent_router_api_key']}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": "You are a professional copy editor. Return only corrected manuscript text.",
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.3,
        }

        try:
            response = requests.post(url, headers=headers, json=payload, timeout=75)
            if response.status_code == 200:
                self._last_ai_warning = ""
                return self._extract_corrected_text(self._extract_openrouter_text(response.json()))
            if response.status_code in (401, 403):
                self._warn_once("AgentRouter request unauthorized/forbidden; check token and model access.")
            elif response.status_code == 429:
                self._warn_once("AgentRouter rate-limited (429); falling back to rule-based editing.")
            else:
                self._warn_once(
                    f"AgentRouter editing failed ({response.status_code}); falling back to rule-based editing."
                )
        except Exception as e:
            self._warn_once(f"AgentRouter editing failed: {e}")

        return None

    def _extract_invariant_tokens(self, text: str) -> Dict[str, Set[str]]:
        """Extract tokens that should remain intact across editing passes."""
        urls = set(re.findall(r'(?i)\b(?:https?|ftp)://[^\s<>"\']+', text))
        urls.update(re.findall(r'(?i)\bwww\.[^\s<>"\']+', text))
        dois = set(re.findall(r'(?i)\bdoi:\s*10\.\d{4,9}/[^\s<>"\']+', text))
        dois.update(re.findall(r'(?i)\b10\.\d{4,9}/[^\s<>"\']+', text))
        emails = set(re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', text))

        citation_numbers: Set[str] = set()
        for group in re.findall(r'\[\s*(\d+(?:\s*,\s*\d+)*)\s*\]', text):
            for part in group.split(','):
                token = part.strip()
                if token:
                    citation_numbers.add(token)

        return {
            "urls": urls,
            "dois": dois,
            "emails": emails,
            "citations": citation_numbers,
        }

    def _candidate_risk_score(self, original: str, candidate: str) -> Tuple[int, List[str]]:
        """Score candidate stability risk; lower is better."""
        penalties = 0
        reasons: List[str] = []

        orig = (original or "").strip()
        cand = (candidate or "").strip()
        if not cand:
            return 999, ["empty output"]

        orig_len = max(1, len(orig))
        cand_len = len(cand)
        len_ratio = cand_len / orig_len
        if len_ratio < 0.60 or len_ratio > 1.60:
            penalties += 35
            reasons.append("large length drift")
        elif len_ratio < 0.75 or len_ratio > 1.35:
            penalties += 12
            reasons.append("moderate length drift")

        orig_paras = max(1, orig.count('\n') + 1)
        cand_paras = max(1, cand.count('\n') + 1)
        para_ratio = cand_paras / orig_paras
        if para_ratio < 0.55 or para_ratio > 1.80:
            penalties += 25
            reasons.append("paragraph structure drift")

        similarity = difflib.SequenceMatcher(a=orig, b=cand, autojunk=False).ratio()
        rewrite_ratio = 1.0 - similarity
        if rewrite_ratio > 0.70:
            penalties += 30
            reasons.append("heavy rewrite")
        elif rewrite_ratio > 0.55:
            penalties += 12
            reasons.append("moderate rewrite")

        orig_tokens = self._extract_invariant_tokens(orig)
        cand_tokens = self._extract_invariant_tokens(cand)

        if orig_tokens["urls"] - cand_tokens["urls"]:
            penalties += 40
            reasons.append("missing urls")
        if orig_tokens["dois"] - cand_tokens["dois"]:
            penalties += 40
            reasons.append("missing doi")
        if orig_tokens["emails"] - cand_tokens["emails"]:
            penalties += 30
            reasons.append("missing emails")

        orig_cites = orig_tokens["citations"]
        cand_cites = cand_tokens["citations"]
        if orig_cites and len(cand_cites) < max(1, int(len(orig_cites) * 0.70)):
            penalties += 20
            reasons.append("citation loss")

        if re.search(r'^\s*references?\s*:?\s*$', orig, flags=re.IGNORECASE | re.MULTILINE):
            if not re.search(r'^\s*references?\s*:?\s*$', cand, flags=re.IGNORECASE | re.MULTILINE):
                penalties += 20
                reasons.append("missing references heading")

        return penalties, reasons

    def _select_best_correction(
        self,
        original: str,
        rules_corrected: str,
        ai_corrected: str,
        prefer_ai: bool = False,
    ) -> str:
        """Select the safest correction output between rule-only and AI-enhanced variants."""
        if not ai_corrected.strip():
            self._last_selection_note = "AI output empty; using rule-based correction."
            summary = self._last_processing_audit.setdefault("summary", {})
            summary["final_selection"] = {
                "decision": "rule_fallback",
                "reason": "empty ai output",
            }
            return rules_corrected

        rule_score, rule_reasons = self._candidate_risk_score(original, rules_corrected)
        ai_score, ai_reasons = self._candidate_risk_score(original, ai_corrected)
        summary = self._last_processing_audit.setdefault("summary", {})

        tolerance = 6 if prefer_ai else 3
        hard_risk_limit = 12 if prefer_ai else None

        # AI candidate must be at least as stable as rule-based output, with a
        # slightly wider lane in AI-first mode as long as structure-risk stays low.
        if ai_score <= rule_score + tolerance and (hard_risk_limit is None or ai_score <= hard_risk_limit):
            decision_label = "AI-first accepted" if prefer_ai else "AI accepted"
            self._last_selection_note = f"{decision_label} (rule_score={rule_score}, ai_score={ai_score})."
            summary["final_selection"] = {
                "decision": "ai_accepted",
                "rule_score": rule_score,
                "ai_score": ai_score,
                "rule_reasons": rule_reasons,
                "ai_reasons": ai_reasons,
                "tolerance": tolerance,
                "prefer_ai": bool(prefer_ai),
                "hard_risk_limit": hard_risk_limit,
            }
            return ai_corrected

        self._warn_once(
            "AI output looked unstable; using safer correction set. "
            f"(rule_score={rule_score}, ai_score={ai_score})"
        )
        self._last_selection_note = (
            f"Rule fallback: rule={rule_score} ({', '.join(rule_reasons) or 'ok'}), "
            f"ai={ai_score} ({', '.join(ai_reasons) or 'ok'})"
        )
        summary["final_selection"] = {
            "decision": "rule_fallback",
            "rule_score": rule_score,
            "ai_score": ai_score,
            "rule_reasons": rule_reasons,
            "ai_reasons": ai_reasons,
            "tolerance": tolerance,
            "prefer_ai": bool(prefer_ai),
            "hard_risk_limit": hard_risk_limit,
        }
        return rules_corrected

    def _warn_once(self, message: str):
        """Print non-fatal AI warning once until state changes."""
        if message != self._last_ai_warning:
            print(message)
            self._last_ai_warning = message

    def _extract_gemini_text(self, response_json: Dict) -> str:
        """Extract response text from Gemini payload."""
        candidates = response_json.get("candidates", [])
        for candidate in candidates:
            content = candidate.get("content", {})
            parts = content.get("parts", [])
            texts = [part.get("text", "") for part in parts if isinstance(part, dict)]
            joined = "".join(texts).strip()
            if joined:
                return joined
        return ""

    def _extract_openrouter_text(self, response_json: Dict) -> str:
        """Extract assistant text from OpenRouter chat completion payload."""
        choices = response_json.get("choices", [])
        for choice in choices:
            message = choice.get("message", {}) if isinstance(choice, dict) else {}
            content = message.get("content", "")
            if isinstance(content, str):
                text = content.strip()
                if text:
                    return text
            elif isinstance(content, list):
                parts: List[str] = []
                for item in content:
                    if isinstance(item, dict):
                        piece = item.get("text", "")
                        if isinstance(piece, str) and piece:
                            parts.append(piece)
                joined = "".join(parts).strip()
                if joined:
                    return joined
        return ""

    def _check_ollama(self, ollama_host: str) -> bool:
        """Check if Ollama is running."""
        try:
            response = requests.get(f"{ollama_host}/api/tags", timeout=5)
            return response.status_code == 200
        except:
            return False

    def _get_ollama_models(self, ollama_host: str) -> List[str]:
        """Return available Ollama model names."""
        names: List[str] = []
        try:
            response = requests.get(f"{ollama_host}/api/tags", timeout=5)
            if response.status_code == 200:
                data = response.json()
                models = data.get("models", []) if isinstance(data, dict) else []
                for model in models:
                    if isinstance(model, dict):
                        name = str(model.get("name", "")).strip()
                        if name and name not in names:
                            names.append(name)
        except Exception:
            pass

        if names:
            return names

        # Fallback for environments where API endpoint isn't reachable but CLI works.
        try:
            proc = subprocess.run(
                ["ollama", "list"],
                check=False,
                capture_output=True,
                text=True,
                timeout=5
            )
            if proc.returncode == 0 and proc.stdout:
                for line in proc.stdout.splitlines()[1:]:
                    line = line.strip()
                    if not line:
                        continue
                    name = line.split()[0].strip()
                    if name and name not in names:
                        names.append(name)
        except Exception:
            pass

        return names

    def _resolve_ollama_model(self, ollama_host: str, requested_model: str) -> Optional[str]:
        """Resolve a usable Ollama model, preferring requested model when available."""
        available = self._get_ollama_models(ollama_host)
        if not available:
            return None

        requested = (requested_model or "").strip()
        if requested:
            if requested in available:
                return requested
            if ":" not in requested:
                for model in available:
                    if model.split(":", 1)[0] == requested:
                        return model

        preferred_bases = ("llama3.1", "llama3", "gemma4", "mistral", "qwen2.5", "llama3.2")
        for base in preferred_bases:
            for model in available:
                if model.split(":", 1)[0] == base:
                    return model

        return available[0]

    def _build_edit_prompt(
        self,
        text: str,
        rules_corrected: str,
        options: Dict,
        stage: str = "full",
        section_index: int = 0,
        section_total: int = 0
    ) -> str:
        """Build the AI editing prompt."""
        instructions = []
        requested_domain = str(options.get('domain_profile', 'auto')).strip().lower()
        detected_domain = self.get_domain_report().get("profile", "general")
        domain_for_prompt = detected_domain if requested_domain == "auto" else requested_domain
        if domain_for_prompt not in ("general", "medical", "engineering", "law"):
            domain_for_prompt = "general"
        journal_profile = self.editor.resolve_journal_profile(options)
        ai_first_cmos = self._is_ai_first_cmos_mode(options)
        initials_rule = (
            "use author initials without periods (e.g., Smith AB)"
            if not bool(journal_profile.get("initials_with_periods"))
            else "use author initials with periods (e.g., Smith A.B.)"
        )
        title_rule = (
            "use sentence case for article/chapter titles"
            if str(journal_profile.get("title_case", "sentence")) == "sentence"
            else "use title case for article/chapter titles"
        )
        journal_rule = (
            "abbreviate journal titles in NLM style"
            if str(journal_profile.get("journal_abbrev", "nlm")) == "nlm"
            else "keep full journal titles (no abbreviation)"
        )

        if options.get('spelling', True):
            instructions.append("- Fix any spelling errors using American spelling (Chicago preference)")
            instructions.append(f"- Preserve {domain_for_prompt} terminology; do not replace technical terms with generic synonyms")
        if options.get('sentence_case', True):
            instructions.append("- Fix capitalization: capitalize first word of sentences, days, months, proper nouns")
        if options.get('punctuation', True):
            instructions.append("- Fix punctuation: proper spacing, quotation marks, ellipsis (...)")
        if options.get('chicago_style', True):
            instructions.append(f"""
- Chicago Manual of Style rules:
  * Use American spelling (ize not ise)
  * Use serial comma
  * Keep author affiliation markers as superscripts (for example: ¹ Dr. Name)
  * Keep in-text citations in numeric bracket format like [9, 19] (no author-year citation text in body)
  * Format the References section in Vancouver numbered style (e.g., [1], [2], [3])
  * Ensure citation numbers and reference numbers follow Vancouver first-appearance order
  * Follow journal profile: {journal_profile.get('label', 'Vancouver')}
  * In references, {initials_rule}
  * In references, {title_rule}
  * In references, {journal_rule}
  * Do NOT alter URLs, DOIs, or email addresses (keep exact spelling, case, and spacing)
  * Italicize titles of complete works
  * Use double quotation marks for quoted material
  * Use em-dash (—) without spaces for interruptions
  * Keep keyword line in this style: Keyword: Key one, Key two, Key three
""")
        if ai_first_cmos:
            instructions.append(
                "- Work as a professional CMOS copy editor: use context-aware grammar, usage, and capitalization judgment instead of rigid pattern substitutions"
            )
            instructions.append(
                "- Preserve meaning and technical claims; improve clarity and correctness without flattening the author's voice"
            )
            instructions.append(
                "- Use the baseline draft as a safety reference only; do not repeat baseline mistakes if your professional judgment indicates a better CMOS correction"
            )

        stage_instruction = ""
        if stage == "section" and section_total > 1:
            stage_instruction = (
                f"\nThis is section {section_index} of {section_total} from a larger manuscript.\n"
                "Edit only this section and preserve line/paragraph boundaries.\n"
            )

        prompt = f"""You are a professional copy editor following The Chicago Manual of Style.
Edit the manuscript and return ONLY the corrected text.
Do NOT add explanations or comments - just the corrected text.
Apply minimal edits only; do not rewrite unchanged sentences.
Preserve structure, references order, citation positions, URLs, DOIs, and emails.
{stage_instruction}

Corrections to apply:
{chr(10).join(instructions)}

Original manuscript:
{text}

Baseline corrected draft (already rule-checked):
{rules_corrected}

Corrected manuscript:"""

        return prompt

    def _extract_corrected_text(self, ai_response: str) -> str:
        """Extract just the corrected text from AI response."""
        # Remove any leading/trailing whitespace
        text = ai_response.strip()

        # Remove common wrapper phrases
        phrases_to_remove = [
            "corrected manuscript:",
            "here's the corrected",
            "corrected text:",
            "edited manuscript:",
        ]

        for phrase in phrases_to_remove:
            if text.lower().startswith(phrase):
                text = text[len(phrase):].strip()

        return text

    def _iter_block_items(self, parent):
        """Yield paragraphs and tables in document order."""
        if isinstance(parent, DocxDocument):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            return

        for child in parent_elm.iterchildren():
            tag = child.tag.rsplit('}', 1)[-1]
            if tag == 'p':
                yield Paragraph(child, parent)
            elif tag == 'tbl':
                yield Table(child, parent)

    def _extract_docx_blocks(self, doc: DocxDocument) -> List[Dict]:
        """Return ordered paragraph/table-row blocks from a DOCX document."""
        blocks: List[Dict] = []
        for block in self._iter_block_items(doc):
            if isinstance(block, Paragraph):
                text_value = self._paragraph_text_for_pipeline(block)
                has_drawing = self._paragraph_has_drawing(block)
                blocks.append({
                    "type": "paragraph",
                    "paragraph": block,
                    "text": text_value,
                    "has_drawing": has_drawing,
                    "consumes_text": bool(text_value.strip()) or not has_drawing,
                })
                continue

            if isinstance(block, Table):
                for row in block.rows:
                    cells = [self._cell_text_for_pipeline(cell) for cell in row.cells]
                    blocks.append({
                        "type": "table_row",
                        "row": row,
                        "cells": cells,
                        "text": '\t'.join(cells),
                        "consumes_text": True,
                    })
        return blocks

    def _paragraph_has_drawing(self, paragraph: Paragraph) -> bool:
        """Return True if paragraph contains drawing/picture elements."""
        try:
            matches = paragraph._p.xpath('.//*[local-name()="drawing" or local-name()="pict"]')
        except Exception:
            matches = []
        return bool(matches)

    def _iter_textbox_containers(self, paragraph: Paragraph) -> List:
        """Return textbox content containers nested inside a paragraph's drawing XML."""
        try:
            return list(paragraph._p.xpath('.//*[local-name()="txbxContent"]'))
        except Exception:
            return []

    def _inspect_docx_package(self, path: str) -> Dict:
        """Inspect special DOCX package parts that require preservation-aware export."""
        report = {
            "comments": 0,
            "footnotes": 0,
            "endnotes": 0,
            "textboxes": 0,
            "preservation_mode": "template_copy_required",
        }
        if not path or not os.path.exists(path):
            return report

        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        def _count_xml_nodes(xml_bytes: bytes, xpath: str) -> int:
            try:
                root = ET.fromstring(xml_bytes)
            except ET.ParseError:
                return 0
            return len(root.findall(xpath, namespace))

        try:
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
                if "word/comments.xml" in names:
                    report["comments"] = _count_xml_nodes(archive.read("word/comments.xml"), ".//w:comment")
                if "word/footnotes.xml" in names:
                    count = _count_xml_nodes(archive.read("word/footnotes.xml"), ".//w:footnote")
                    report["footnotes"] = max(0, count - 2)
                if "word/endnotes.xml" in names:
                    count = _count_xml_nodes(archive.read("word/endnotes.xml"), ".//w:endnote")
                    report["endnotes"] = max(0, count - 2)
                if "word/document.xml" in names:
                    report["textboxes"] = _count_xml_nodes(archive.read("word/document.xml"), './/w:txbxContent')
        except Exception:
            return report

        return report

    def _merge_content_types_override(self, root, part_name: str, content_type: str):
        """Ensure a content-types Override exists for the given part."""
        namespace = "http://schemas.openxmlformats.org/package/2006/content-types"
        override_tag = f"{{{namespace}}}Override"
        for node in root.findall(override_tag):
            if node.attrib.get("PartName") == part_name:
                return
        ET.SubElement(root, override_tag, {"PartName": part_name, "ContentType": content_type})

    def _merge_relationships_from_source(self, source_root, output_root, relationship_types: Set[str]):
        """Merge document relationship entries from source into output."""
        namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
        rel_tag = f"{{{namespace}}}Relationship"
        existing_targets = {
            (node.attrib.get("Type"), node.attrib.get("Target"))
            for node in output_root.findall(rel_tag)
        }
        next_id = 1
        for node in output_root.findall(rel_tag):
            raw_id = str(node.attrib.get("Id", "") or "")
            if raw_id.startswith("rId"):
                try:
                    next_id = max(next_id, int(raw_id[3:]) + 1)
                except ValueError:
                    continue
        for node in source_root.findall(rel_tag):
            rel_type = node.attrib.get("Type")
            target = node.attrib.get("Target")
            if rel_type not in relationship_types or not target:
                continue
            key = (rel_type, target)
            if key in existing_targets:
                continue
            attrs = dict(node.attrib)
            attrs["Id"] = f"rId{next_id}"
            next_id += 1
            ET.SubElement(output_root, rel_tag, attrs)
            existing_targets.add(key)

    def _preserve_docx_special_parts(self, source_docx_path: str, output_path: str):
        """Copy comments/footnotes/endnotes package parts from source to output."""
        if not source_docx_path or not output_path:
            return
        if not (os.path.exists(source_docx_path) and os.path.exists(output_path)):
            return

        special_parts = {
            "word/comments.xml": "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
            "word/footnotes.xml": "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
            "word/endnotes.xml": "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml",
        }
        relationship_types = {
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
        }

        with zipfile.ZipFile(source_docx_path, "r") as source_zip, zipfile.ZipFile(output_path, "r") as output_zip:
            source_names = set(source_zip.namelist())
            output_names = list(output_zip.namelist())
            replacement_parts = {part_name for part_name in special_parts if part_name in source_names}
            replacement_parts.update({"[Content_Types].xml", "word/_rels/document.xml.rels"})

            rebuilt_entries: Dict[str, bytes] = {}
            for name in output_names:
                if name in rebuilt_entries:
                    continue
                if name in replacement_parts:
                    continue
                rebuilt_entries[name] = output_zip.read(name)

            for part_name in special_parts:
                if part_name in source_names:
                    rebuilt_entries[part_name] = source_zip.read(part_name)

            if "[Content_Types].xml" in source_names and "[Content_Types].xml" in output_names:
                output_root = ET.fromstring(output_zip.read("[Content_Types].xml"))
                for part_name, content_type in special_parts.items():
                    if part_name in source_names:
                        self._merge_content_types_override(output_root, "/" + part_name, content_type)
                rebuilt_entries["[Content_Types].xml"] = ET.tostring(output_root, encoding="utf-8", xml_declaration=True)

            rels_name = "word/_rels/document.xml.rels"
            if rels_name in source_names and rels_name in output_names:
                source_rels = ET.fromstring(source_zip.read(rels_name))
                output_rels = ET.fromstring(output_zip.read(rels_name))
                self._merge_relationships_from_source(source_rels, output_rels, relationship_types)
                rebuilt_entries[rels_name] = ET.tostring(output_rels, encoding="utf-8", xml_declaration=True)

        output_dir = os.path.dirname(os.path.abspath(output_path)) or "."
        os.makedirs(output_dir, exist_ok=True)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=output_dir) as temp_zip_handle:
            temp_zip_path = temp_zip_handle.name
        try:
            with zipfile.ZipFile(temp_zip_path, "w", compression=zipfile.ZIP_DEFLATED) as rebuilt_zip:
                for name, payload in rebuilt_entries.items():
                    rebuilt_zip.writestr(name, payload)
            os.replace(temp_zip_path, output_path)
        finally:
            if os.path.exists(temp_zip_path):
                os.unlink(temp_zip_path)

    def _textbox_container_text(self, container) -> str:
        """Serialize textbox content into a plain-text pipeline string."""
        values: List[str] = []
        for child in list(container):
            if child.tag.rsplit('}', 1)[-1] != 'p':
                continue
            paragraph = Paragraph(child, None)
            values.append(paragraph.text or "")
        if not values:
            return ""
        return self.TEXTBOX_PARAGRAPH_MARKER.join(values)

    def _paragraph_textbox_texts(self, paragraph: Paragraph) -> List[str]:
        """Return serialized text for each textbox found in a paragraph."""
        return [self._textbox_container_text(container) for container in self._iter_textbox_containers(paragraph)]

    def _paragraph_text_for_pipeline(self, paragraph: Paragraph) -> str:
        """Serialize paragraph body text and textbox text for processing."""
        body_text = paragraph.text or ""
        textbox_texts = self._paragraph_textbox_texts(paragraph)
        if not textbox_texts:
            return body_text
        if body_text.strip():
            return self.TEXTBOX_MARKER.join([body_text] + textbox_texts)
        if len(textbox_texts) == 1:
            return textbox_texts[0]
        return self.TEXTBOX_MARKER.join(textbox_texts)

    def _clear_paragraph_content(self, paragraph: Paragraph, keep_drawings: bool = False):
        """Remove paragraph content while optionally preserving image runs."""
        p = paragraph._p
        for child in list(p):
            if child.tag.rsplit('}', 1)[-1] == 'pPr':
                continue
            if keep_drawings:
                try:
                    if child.xpath('.//*[local-name()="drawing" or local-name()="pict"]'):
                        continue
                except Exception:
                    pass
            p.remove(child)

    def _remove_paragraph(self, paragraph: Paragraph):
        """Delete paragraph from document tree."""
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    def _remove_table(self, table: Table):
        """Delete table from document tree."""
        element = table._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    def _append_text_segments_to_paragraph(self, paragraph: Paragraph, text: str, segment_type: Optional[str] = None):
        """Append styled text to an existing paragraph."""
        if not text:
            return
        for is_missing, segment in self._iter_missing_placeholder_segments(text):
            if is_missing:
                self._append_docx_run(paragraph, segment, segment_type=segment_type, is_missing=True)
                continue
            for is_foreign, foreign_segment in self._iter_foreign_segments(segment):
                self._append_docx_run(paragraph, foreign_segment, segment_type=segment_type, is_foreign=is_foreign)

    def _configure_docx_document_defaults(self, doc: DocxDocument):
        """Apply base page and font defaults for generated DOCX documents."""
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    def _resolve_docx_style_name(self, doc: DocxDocument, preferred: str, fallback: str = "Normal") -> str:
        """Return a valid style name present in the document."""
        try:
            doc.styles[preferred]
            return preferred
        except KeyError:
            return fallback

    def _classify_plaintext_line(self, line: str) -> Dict[str, str]:
        """Infer a DOCX paragraph role from a plain-text line."""
        raw = str(line or "")
        stripped = raw.strip()
        if not stripped:
            return {"kind": "blank", "text": ""}

        heading_match = self.HEADING_LINE_RE.match(raw)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            return {"kind": f"heading_{level}", "text": heading_match.group(2).strip()}

        bullet_match = self.BULLET_LINE_RE.match(raw)
        if bullet_match:
            return {"kind": "bullet", "text": bullet_match.group(2).strip()}

        numbered_match = self.NUMBERED_LINE_RE.match(raw)
        if numbered_match:
            return {"kind": "number", "text": numbered_match.group(2).strip()}

        if (
            self.PLAINTEXT_HEADING_RE.match(stripped)
            and '\t' not in stripped
            and len(stripped.split()) <= 8
            and stripped[-1] not in '.!?'
        ):
            return {"kind": "heading_1", "text": stripped.rstrip(':')}

        return {"kind": "body", "text": raw}

    def _apply_fallback_paragraph_layout(self, paragraph: Paragraph, kind: str):
        """Apply paragraph style/spacing for fallback DOCX generation."""
        doc = paragraph.part.document
        style_name = "Normal"
        if kind == "bullet":
            style_name = self._resolve_docx_style_name(doc, "List Bullet")
        elif kind == "number":
            style_name = self._resolve_docx_style_name(doc, "List Number")
        elif kind.startswith("heading_"):
            level = kind.rsplit('_', 1)[-1]
            style_name = self._resolve_docx_style_name(doc, f"Heading {level}")

        paragraph.style = doc.styles[style_name]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5

        if kind.startswith("heading_"):
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
        elif kind == "blank":
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    def _add_fallback_docx_paragraph(
        self,
        doc: DocxDocument,
        text: str,
        *,
        highlighted: bool = False,
        original_text: str = "",
        structure_hint: Optional[str] = None,
    ):
        """Append one inferred paragraph to a generated fallback DOCX."""
        classified = self._classify_plaintext_line(structure_hint if structure_hint is not None else text)
        paragraph = doc.add_paragraph()
        self._apply_fallback_paragraph_layout(paragraph, classified["kind"])

        body_text = self._classify_plaintext_line(text)["text"]
        if highlighted:
            original_body_text = self._classify_plaintext_line(original_text)["text"]
            for segment_type, segment_text in self._iter_diff_segments(original_body_text, body_text):
                self._append_text_segments_to_paragraph(paragraph, segment_text, segment_type=segment_type)
        else:
            self._append_text_segments_to_paragraph(paragraph, body_text)
        return paragraph

    def _split_textbox_paragraphs(self, text: str) -> List[str]:
        """Deserialize textbox content into paragraph chunks."""
        raw = str(text or "")
        if self.TEXTBOX_PARAGRAPH_MARKER not in raw:
            return [raw]
        return raw.split(self.TEXTBOX_PARAGRAPH_MARKER)

    def _split_paragraph_and_textboxes(self, paragraph: Paragraph, text: str) -> Tuple[str, List[str]]:
        """Split serialized paragraph pipeline text into body and textbox segments."""
        raw = str(text or "")
        current_textboxes = self._paragraph_textbox_texts(paragraph)
        if not current_textboxes:
            return raw, []

        if self.TEXTBOX_MARKER in raw:
            parts = raw.split(self.TEXTBOX_MARKER)
            if len(parts) >= len(current_textboxes) + 1:
                body_text = parts[0]
                textbox_texts = parts[1 : len(current_textboxes) + 1]
                if len(textbox_texts) < len(current_textboxes):
                    textbox_texts.extend(current_textboxes[len(textbox_texts) :])
                return body_text, textbox_texts

        if (paragraph.text or "").strip():
            return raw, current_textboxes

        if len(current_textboxes) == 1:
            return "", [raw]

        textbox_texts = [raw] + current_textboxes[1:]
        return "", textbox_texts

    def _sync_textbox_container(self, paragraph: Paragraph, container, text: str, original_text: str = "", highlighted: bool = False):
        """Rewrite textbox XML paragraphs while preserving the shape container."""
        desired_parts = self._split_textbox_paragraphs(text)
        original_parts = self._split_textbox_paragraphs(original_text)
        desired_count = max(len(desired_parts), 1)

        paragraphs = [child for child in list(container) if child.tag.rsplit('}', 1)[-1] == 'p']
        while len(paragraphs) < desired_count:
            new_paragraph = OxmlElement('w:p')
            container.append(new_paragraph)
            paragraphs.append(new_paragraph)
        while len(paragraphs) > desired_count:
            container.remove(paragraphs.pop())

        for idx, paragraph_elm in enumerate(paragraphs):
            textbox_paragraph = Paragraph(paragraph_elm, paragraph._parent)
            current_original = original_parts[idx] if idx < len(original_parts) else ""
            current_desired = desired_parts[idx] if idx < len(desired_parts) else ""
            self._clear_paragraph_content(textbox_paragraph)
            if highlighted:
                for segment_type, segment_text in self._iter_diff_segments(current_original, current_desired):
                    self._append_text_segments_to_paragraph(textbox_paragraph, segment_text, segment_type=segment_type)
            else:
                self._append_text_segments_to_paragraph(textbox_paragraph, current_desired)

    def _sync_textboxes(self, paragraph: Paragraph, text: str, original_text: str = "", highlighted: bool = False):
        """Rewrite textbox contents referenced by a paragraph's drawing XML."""
        containers = self._iter_textbox_containers(paragraph)
        if not containers:
            return

        _, desired_textboxes = self._split_paragraph_and_textboxes(paragraph, text)
        _, original_textboxes = self._split_paragraph_and_textboxes(paragraph, original_text)
        if not desired_textboxes:
            desired_textboxes = self._paragraph_textbox_texts(paragraph)
        if not original_textboxes:
            original_textboxes = self._paragraph_textbox_texts(paragraph)

        for idx, container in enumerate(containers):
            desired_value = desired_textboxes[idx] if idx < len(desired_textboxes) else ""
            original_value = original_textboxes[idx] if idx < len(original_textboxes) else ""
            self._sync_textbox_container(
                paragraph,
                container,
                desired_value,
                original_text=original_value,
                highlighted=highlighted,
            )

    def _write_paragraph_text(self, paragraph: Paragraph, text: str):
        """Replace paragraph text while preserving paragraph formatting and drawings."""
        keep_drawings = self._paragraph_has_drawing(paragraph)
        body_text, _ = self._split_paragraph_and_textboxes(paragraph, text)
        self._clear_paragraph_content(paragraph, keep_drawings=keep_drawings)
        self._append_text_segments_to_paragraph(paragraph, body_text)
        self._sync_textboxes(paragraph, text, highlighted=False)

    def _write_paragraph_diff(self, paragraph: Paragraph, original: str, corrected: str):
        """Replace paragraph content with redline-style runs while preserving drawings."""
        keep_drawings = self._paragraph_has_drawing(paragraph)
        original_body, _ = self._split_paragraph_and_textboxes(paragraph, original)
        corrected_body, _ = self._split_paragraph_and_textboxes(paragraph, corrected)
        self._clear_paragraph_content(paragraph, keep_drawings=keep_drawings)
        for segment_type, segment_text in self._iter_diff_segments(original_body, corrected_body):
            self._append_text_segments_to_paragraph(paragraph, segment_text, segment_type=segment_type)
        self._sync_textboxes(paragraph, corrected, original_text=original, highlighted=True)

    def _split_table_line(self, line: str, cell_count: int) -> List[str]:
        """Split a flat table-row line back into cell values."""
        if cell_count <= 1:
            return [line or ""]
        values = (line or "").split('\t')
        if len(values) < cell_count:
            values.extend([""] * (cell_count - len(values)))
        elif len(values) > cell_count:
            values = values[: cell_count - 1] + ['\t'.join(values[cell_count - 1 :])]
        return values

    def _cell_text_for_pipeline(self, cell) -> str:
        """Serialize cell text for the text-processing pipeline."""
        blocks: List[str] = []
        contains_nested_table = False
        cell_blocks = list(self._iter_block_items(cell))

        for index, block in enumerate(cell_blocks):
            if isinstance(block, Paragraph):
                has_neighboring_table = (
                    (index > 0 and isinstance(cell_blocks[index - 1], Table))
                    or (index + 1 < len(cell_blocks) and isinstance(cell_blocks[index + 1], Table))
                )
                if not (block.text or "").strip() and not self._paragraph_has_drawing(block) and has_neighboring_table:
                    continue
                blocks.append(f"P:{block.text or ''}")
                continue

            contains_nested_table = True
            row_values: List[str] = []
            for row in block.rows:
                nested_cells = [self._cell_text_for_pipeline(nested_cell) for nested_cell in row.cells]
                row_values.append(self.CELL_TABLE_CELL_MARKER.join(nested_cells))
            blocks.append(f"T:{self.CELL_TABLE_ROW_MARKER.join(row_values)}")

        if not blocks:
            return ""

        if not contains_nested_table:
            paragraph_values = [block[2:] for block in blocks if block.startswith("P:")]
            return self.CELL_PARAGRAPH_MARKER.join(paragraph_values)

        return self.CELL_BLOCK_MARKER.join(blocks)

    def _split_cell_paragraphs(self, text: str) -> List[str]:
        """Deserialize pipeline cell text into paragraph chunks."""
        raw = str(text or "")
        if self.CELL_PARAGRAPH_MARKER not in raw:
            return [raw]
        return raw.split(self.CELL_PARAGRAPH_MARKER)

    def _split_cell_blocks(self, text: str) -> List[Dict]:
        """Deserialize pipeline cell text into paragraph/table blocks."""
        raw = str(text or "")
        if self.CELL_BLOCK_MARKER not in raw and not raw.startswith(("P:", "T:")):
            return [{"type": "paragraph", "text": part} for part in self._split_cell_paragraphs(raw)]

        segments = raw.split(self.CELL_BLOCK_MARKER)
        blocks: List[Dict] = []
        for segment in segments:
            if segment.startswith("T:"):
                payload = segment[2:]
                rows = payload.split(self.CELL_TABLE_ROW_MARKER) if payload else []
                parsed_rows: List[List[str]] = []
                for row in rows:
                    parsed_rows.append(row.split(self.CELL_TABLE_CELL_MARKER) if row else [""])
                blocks.append({"type": "table", "rows": parsed_rows})
                continue

            payload = segment[2:] if segment.startswith("P:") else segment
            for part in self._split_cell_paragraphs(payload):
                blocks.append({"type": "paragraph", "text": part})
        return blocks

    def _rewrite_table(self, table: Table, rows: List[List[str]]):
        """Rewrite a table and nested cell content using serialized row values."""
        for row_index, row in enumerate(table.rows):
            row_values = rows[row_index] if row_index < len(rows) else []
            for cell_index, cell in enumerate(row.cells):
                corrected_value = row_values[cell_index] if cell_index < len(row_values) else ""
                self._rewrite_cell(cell, corrected_value)

    def _rewrite_table_diff(self, table: Table, original_rows: List[List[str]], corrected_rows: List[List[str]]):
        """Rewrite a table with diff styling using serialized row values."""
        for row_index, row in enumerate(table.rows):
            original_row = original_rows[row_index] if row_index < len(original_rows) else []
            corrected_row = corrected_rows[row_index] if row_index < len(corrected_rows) else []
            for cell_index, cell in enumerate(row.cells):
                original_value = original_row[cell_index] if cell_index < len(original_row) else ""
                corrected_value = corrected_row[cell_index] if cell_index < len(corrected_row) else ""
                self._rewrite_cell_diff(cell, original_value, corrected_value)

    def _rewrite_cell(self, cell, text: str):
        """Replace all visible text inside a table cell."""
        parsed_blocks = self._split_cell_blocks(text)
        paragraph_blocks = [block for block in parsed_blocks if block["type"] == "paragraph"]
        table_blocks = [block for block in parsed_blocks if block["type"] == "table"]

        existing_blocks = list(self._iter_block_items(cell))
        paragraph_index = 0
        table_index = 0

        for block_position, block in enumerate(existing_blocks):
            if isinstance(block, Paragraph):
                remaining_existing_paragraphs = sum(
                    1 for later in existing_blocks[block_position:] if isinstance(later, Paragraph)
                )
                remaining_paragraphs = len(paragraph_blocks) - paragraph_index
                if (
                    not (block.text or "").strip()
                    and not self._paragraph_has_drawing(block)
                    and remaining_existing_paragraphs > remaining_paragraphs
                ):
                    self._remove_paragraph(block)
                    continue
                block_text = paragraph_blocks[paragraph_index]["text"] if paragraph_index < len(paragraph_blocks) else ""
                paragraph_index += 1
                self._clear_paragraph_content(block)
                self._append_text_segments_to_paragraph(block, block_text)
                continue

            if table_index < len(table_blocks):
                self._rewrite_table(block, table_blocks[table_index]["rows"])
                table_index += 1

        for extra_paragraph in paragraph_blocks[paragraph_index:]:
            new_paragraph = cell.add_paragraph()
            self._clear_paragraph_content(new_paragraph)
            self._append_text_segments_to_paragraph(new_paragraph, extra_paragraph["text"])

    def _rewrite_cell_diff(self, cell, original: str, corrected: str):
        """Replace cell text with redline-style runs."""
        original_blocks = self._split_cell_blocks(original)
        corrected_blocks = self._split_cell_blocks(corrected)
        original_paragraphs = [block for block in original_blocks if block["type"] == "paragraph"]
        corrected_paragraphs = [block for block in corrected_blocks if block["type"] == "paragraph"]
        original_tables = [block for block in original_blocks if block["type"] == "table"]
        corrected_tables = [block for block in corrected_blocks if block["type"] == "table"]

        existing_blocks = list(self._iter_block_items(cell))
        paragraph_index = 0
        table_index = 0

        for block_position, block in enumerate(existing_blocks):
            if isinstance(block, Paragraph):
                remaining_existing_paragraphs = sum(
                    1 for later in existing_blocks[block_position:] if isinstance(later, Paragraph)
                )
                remaining_paragraphs = max(len(original_paragraphs), len(corrected_paragraphs)) - paragraph_index
                if (
                    not (block.text or "").strip()
                    and not self._paragraph_has_drawing(block)
                    and remaining_existing_paragraphs > remaining_paragraphs
                ):
                    self._remove_paragraph(block)
                    continue
                original_text = original_paragraphs[paragraph_index]["text"] if paragraph_index < len(original_paragraphs) else ""
                corrected_text = corrected_paragraphs[paragraph_index]["text"] if paragraph_index < len(corrected_paragraphs) else ""
                paragraph_index += 1
                self._clear_paragraph_content(block)
                for segment_type, segment_text in self._iter_diff_segments(original_text, corrected_text):
                    self._append_text_segments_to_paragraph(block, segment_text, segment_type=segment_type)
                continue

            original_rows = original_tables[table_index]["rows"] if table_index < len(original_tables) else []
            corrected_rows = corrected_tables[table_index]["rows"] if table_index < len(corrected_tables) else []
            table_index += 1
            self._rewrite_table_diff(block, original_rows, corrected_rows)

        pair_count = max(len(original_paragraphs), len(corrected_paragraphs))
        for idx in range(paragraph_index, pair_count):
            target_paragraph = cell.add_paragraph()
            self._clear_paragraph_content(target_paragraph)
            original_text = original_paragraphs[idx]["text"] if idx < len(original_paragraphs) else ""
            corrected_text = corrected_paragraphs[idx]["text"] if idx < len(corrected_paragraphs) else ""
            for segment_type, segment_text in self._iter_diff_segments(original_text, corrected_text):
                self._append_text_segments_to_paragraph(target_paragraph, segment_text, segment_type=segment_type)

    def _apply_text_to_template_docx(self, source_docx_path: str, text: str, output_path: str, highlighted: bool = False):
        """Project corrected text back into the original DOCX structure."""
        doc = Document(source_docx_path)
        blocks = self._extract_docx_blocks(doc)
        corrected_lines = (text or "").split('\n')
        corrected_index = 0

        for block in blocks:
            if block.get("consumes_text", True):
                corrected_line = corrected_lines[corrected_index] if corrected_index < len(corrected_lines) else ""
                corrected_index += 1
            else:
                corrected_line = block.get("text", "")
            if block["type"] == "paragraph":
                if highlighted:
                    self._write_paragraph_diff(block["paragraph"], block["text"], corrected_line)
                else:
                    self._write_paragraph_text(block["paragraph"], corrected_line)
                continue

            cell_values = self._split_table_line(corrected_line, len(block["cells"]))
            for cell, original_value, corrected_value in zip(block["row"].cells, block["cells"], cell_values):
                if highlighted:
                    self._rewrite_cell_diff(cell, original_value, corrected_value)
                else:
                    self._rewrite_cell(cell, corrected_value)

        for extra_line in corrected_lines[corrected_index:]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            if highlighted:
                self._write_paragraph_diff(paragraph, "", extra_line)
            else:
                self._write_paragraph_text(paragraph, extra_line)

        doc.save(output_path)
        self._preserve_docx_special_parts(source_docx_path, output_path)

    def generate_clean_docx(self, text: str, output_path: str, source_docx_path: str = ""):
        """Generate a clean DOCX with all corrections applied."""
        if source_docx_path and os.path.exists(source_docx_path):
            self._apply_text_to_template_docx(source_docx_path, text, output_path, highlighted=False)
            return

        doc = Document()
        self._configure_docx_document_defaults(doc)

        for paragraph_text in text.split('\n'):
            self._add_fallback_docx_paragraph(doc, paragraph_text, highlighted=False)

        doc.save(output_path)

    def generate_highlighted_docx(self, original: str, corrected: str, output_path: str, source_docx_path: str = ""):
        """Generate a DOCX with track changes showing corrections."""
        if source_docx_path and os.path.exists(source_docx_path):
            self._apply_text_to_template_docx(source_docx_path, corrected, output_path, highlighted=True)
            return

        doc = Document()
        self._configure_docx_document_defaults(doc)

        original_lines = (original or "").split('\n')
        corrected_lines = (corrected or "").split('\n')
        matcher = difflib.SequenceMatcher(a=original_lines, b=corrected_lines, autojunk=False)

        for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
            if opcode == "equal":
                for line in corrected_lines[j1:j2]:
                    self._add_fallback_docx_paragraph(doc, line, highlighted=False)
                continue

            if opcode == "insert":
                for line in corrected_lines[j1:j2]:
                    self._add_fallback_docx_paragraph(
                        doc,
                        line,
                        highlighted=True,
                        original_text="",
                    )
                continue

            if opcode == "delete":
                for line in original_lines[i1:i2]:
                    self._add_fallback_docx_paragraph(
                        doc,
                        "",
                        highlighted=True,
                        original_text=line,
                        structure_hint=line,
                    )
                continue

            left = original_lines[i1:i2]
            right = corrected_lines[j1:j2]
            pair_count = max(len(left), len(right))
            for idx in range(pair_count):
                original_line = left[idx] if idx < len(left) else ""
                corrected_line = right[idx] if idx < len(right) else ""
                structure_hint = corrected_line if corrected_line else original_line
                self._add_fallback_docx_paragraph(
                    doc,
                    corrected_line,
                    highlighted=True,
                    original_text=original_line,
                    structure_hint=structure_hint,
                )

        doc.save(output_path)

    def build_redline_html(self, original: str, corrected: str) -> str:
        """Build redline HTML preview with Word-style red change markup."""
        chunks = []
        for segment_type, segment_text in self._iter_diff_segments(original, corrected):
            escaped = self._build_annotated_html(segment_text, include_foreign=False)
            if segment_type == "delete":
                chunks.append(f'<span class="redline-del">{escaped}</span>')
            elif segment_type == "insert":
                chunks.append(f'<span class="redline-add">{escaped}</span>')
            else:
                chunks.append(escaped)
        return "".join(chunks)

    def _foreign_terms_catalog(self) -> List[str]:
        """Return sorted foreign terms that should be italicized in output."""
        if hasattr(self.editor, "get_foreign_term_style_catalog"):
            catalog = self.editor.get_foreign_term_style_catalog() or {}
            terms = catalog.get("italic", []) or []
        else:
            terms = getattr(self.editor, "FOREIGN_TERMS", set()) or set()
        return sorted({str(t).strip().lower() for t in terms if str(t).strip()}, key=len, reverse=True)

    def _protected_literal_spans(self, text: str) -> List[Tuple[int, int]]:
        """Return merged spans for literals that must never be foreign-term styled."""
        source = text or ""
        if not source:
            return []

        patterns = [
            re.compile(r'(?i)\b(?:https?|ftp)://[^\s<>"\']+'),
            re.compile(r'(?i)\bwww\.[^\s<>"\']+'),
            re.compile(r'(?i)\bdoi:\s*10\.\d{4,9}/[^\s<>"\']+'),
            re.compile(r'(?i)\b10\.\d{4,9}/[^\s<>"\']+'),
            re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'),
        ]

        spans: List[Tuple[int, int]] = []
        for pattern in patterns:
            for match in pattern.finditer(source):
                spans.append((match.start(), match.end()))

        if not spans:
            return []

        spans.sort(key=lambda item: (item[0], item[1]))
        merged: List[List[int]] = []
        for start, end in spans:
            if not merged or start > merged[-1][1]:
                merged.append([start, end])
                continue
            merged[-1][1] = max(int(merged[-1][1]), int(end))
        return [(int(start), int(end)) for start, end in merged]

    def _span_overlaps(self, spans: List[Tuple[int, int]], start: int, end: int) -> bool:
        """Return True when [start, end) overlaps any span in a sorted span list."""
        for span_start, span_end in spans:
            if end <= span_start:
                break
            if start < span_end and end > span_start:
                return True
        return False

    def _iter_foreign_segments(self, text: str):
        """Yield tuple (is_foreign, segment_text) by matching known foreign terms."""
        source = text or ""
        terms = self._foreign_terms_catalog()
        if not source or not terms:
            if source:
                yield False, source
            return

        pattern = re.compile(
            r'(?i)(?<!\w)(' + '|'.join(re.escape(term) for term in terms) + r')(?!\w)'
        )
        protected_spans = self._protected_literal_spans(source)

        last = 0
        for match in pattern.finditer(source):
            start, end = match.span()
            if self._span_overlaps(protected_spans, start, end):
                continue
            if match.start() > last:
                yield False, source[last:match.start()]
            yield True, match.group(0)
            last = match.end()

        if last < len(source):
            yield False, source[last:]

    def _iter_missing_placeholder_segments(self, text: str):
        """Yield tuple (is_missing_placeholder, segment_text)."""
        source = text or ""
        pattern = getattr(self.editor, "MISSING_PLACEHOLDER_RE", None)
        if not source or pattern is None:
            if source:
                yield False, source
            return

        last = 0
        for match in pattern.finditer(source):
            if match.start() > last:
                yield False, source[last:match.start()]
            yield True, match.group(0)
            last = match.end()

        if last < len(source):
            yield False, source[last:]

    def _append_docx_run(self, paragraph, text: str, *, segment_type: Optional[str] = None, is_foreign: bool = False, is_missing: bool = False):
        """Append a styled DOCX run for preview/export output."""
        run = paragraph.add_run(text)
        if is_foreign:
            run.italic = True
        if segment_type == "delete":
            run.font.strike = True
        elif segment_type == "insert":
            run.font.underline = True

        if is_missing:
            run.font.color.rgb = self.MISSING_PLACEHOLDER_COLOR
        elif segment_type in ("delete", "insert"):
            run.font.color.rgb = RGBColor(200, 0, 0)
        return run

    def _build_annotated_html(self, text: str, include_foreign: bool = True) -> str:
        """Return HTML-safe text with placeholders/foreign terms wrapped for display."""
        parts: List[str] = []
        for is_missing, segment in self._iter_missing_placeholder_segments(text):
            if is_missing:
                parts.append(f'<span class="missing-placeholder">{html.escape(segment)}</span>')
                continue
            if include_foreign:
                for is_foreign, sub_segment in self._iter_foreign_segments(segment):
                    escaped = html.escape(sub_segment)
                    if is_foreign:
                        parts.append(f'<em class="foreign-term">{escaped}</em>')
                    else:
                        parts.append(escaped)
            else:
                parts.append(html.escape(segment))
        return "".join(parts)

    def build_foreign_annotated_html(self, text: str) -> str:
        """Return HTML-safe text with foreign terms wrapped for italic rendering."""
        lines = (text or "").split('\n')
        out_lines: List[str] = []

        for line in lines:
            out_lines.append(self._build_annotated_html(line, include_foreign=True))

        return '\n'.join(out_lines)

    def get_domain_report(self) -> Dict:
        """Return domain dictionary usage report from editor."""
        if hasattr(self.editor, "get_domain_report"):
            return self.editor.get_domain_report()
        return {
            "profile": "general",
            "scores": {"medical": 0, "engineering": 0, "law": 0},
            "protected_terms": 0,
        }

    def get_processing_audit(self) -> Dict:
        """Return last processing audit details (chunk decisions, scores, summary)."""
        if not self._last_processing_audit:
            return {"mode": "unknown", "sections": [], "summary": {}}
        self._attach_docx_package_summary()
        return self._last_processing_audit

    def get_journal_profile_report(self) -> Dict:
        """Return profile-aware reference formatting audit for last run."""
        if self._last_journal_profile_report:
            return self._last_journal_profile_report
        return self.editor.build_reference_profile_report("", {"journal_profile": "vancouver_periods"})

    def get_citation_reference_report(self) -> Dict:
        """Return citation/reference validator report for last run."""
        if self._last_citation_reference_report:
            return self._last_citation_reference_report
        return self.editor.build_citation_reference_validator_report("", {"journal_profile": "vancouver_periods"})

    def build_noun_report(self, text: str) -> Dict:
        """Identify proper/common nouns, preferring spaCy and falling back to heuristics."""
        content = text or ""
        if not content.strip():
            return {"source": "none", "proper_nouns": [], "common_nouns": []}

        spacy_report = self._build_noun_report_spacy(content)
        if spacy_report is not None:
            return spacy_report

        return self._build_noun_report_heuristic(content)

    def _build_noun_report_spacy(self, text: str) -> Optional[Dict]:
        """Use spaCy POS tags when available."""
        try:
            if self._nlp is None:
                import spacy  # type: ignore
                self._nlp = spacy.load("en_core_web_sm")
            doc = self._nlp(text)
        except Exception:
            return None

        proper_counter: Counter = Counter()
        common_counter: Counter = Counter()

        for token in doc:
            if not token.is_alpha:
                continue
            value = token.text.strip()
            if len(value) < 2:
                continue
            if token.pos_ == "PROPN":
                proper_counter[value] += 1
            elif token.pos_ == "NOUN":
                common_counter[value.lower()] += 1

        return {
            "source": "spacy",
            "proper_nouns": [{"word": k, "count": v} for k, v in proper_counter.most_common(60)],
            "common_nouns": [{"word": k, "count": v} for k, v in common_counter.most_common(80)],
        }

    def _build_noun_report_heuristic(self, text: str) -> Dict:
        """Heuristic noun extraction fallback when spaCy is unavailable."""
        stopwords = {
            'a', 'an', 'and', 'are', 'as', 'at', 'be', 'been', 'but', 'by', 'for', 'from', 'had',
            'has', 'have', 'he', 'her', 'his', 'in', 'is', 'it', 'its', 'of', 'on', 'or', 'she',
            'that', 'the', 'their', 'them', 'there', 'these', 'they', 'this', 'to', 'was', 'were',
            'will', 'with', 'we', 'you', 'your', 'our', 'not', 'no', 'yes'
        }

        proper_counter: Counter = Counter()
        common_counter: Counter = Counter()

        sentence_split = re.split(r'(?<=[.!?])\s+', text)
        for sentence in sentence_split:
            words = re.findall(r"[A-Za-z][A-Za-z'’\-]*", sentence)
            for idx, word in enumerate(words):
                lower = word.lower()
                if len(lower) < 2 or lower in stopwords:
                    continue

                looks_proper = False
                if re.fullmatch(r"[A-Z]{2,}", word):
                    looks_proper = True
                elif word[0].isupper() and idx > 0:
                    looks_proper = True

                if looks_proper:
                    proper_counter[word] += 1
                else:
                    common_counter[lower] += 1

        return {
            "source": "heuristic",
            "proper_nouns": [{"word": k, "count": v} for k, v in proper_counter.most_common(60)],
            "common_nouns": [{"word": k, "count": v} for k, v in common_counter.most_common(80)],
        }

    def build_corrections_report(self, original: str, corrected: str) -> Dict:
        """Build grouped correction list for quick review in UI."""
        group_order = ("spelling", "capitalization", "punctuation", "citation", "reference", "style")
        groups: Dict[str, List[Dict]] = {key: [] for key in group_order}
        seen = set()
        max_total_items = 1200

        original_lines = original.split('\n')
        corrected_lines = corrected.split('\n')
        original_ref_flags = self._reference_line_flags(original_lines)
        corrected_ref_flags = self._reference_line_flags(corrected_lines)

        line_matcher = difflib.SequenceMatcher(a=original_lines, b=corrected_lines, autojunk=False)

        def push_item(kind: str, line_number: int, old_text: str, new_text: str, context_line: str):
            if kind not in groups:
                kind = "style"
            if sum(len(v) for v in groups.values()) >= max_total_items:
                return

            old_text = (old_text or "").strip()
            new_text = (new_text or "").strip()
            context_line = (context_line or "").strip()
            if not old_text and not new_text:
                return

            if len(context_line) > 220:
                context_line = context_line[:217] + "..."

            key = (kind, line_number, old_text, new_text)
            if key in seen:
                return
            seen.add(key)

            groups[kind].append({
                "line": max(1, int(line_number)),
                "original": old_text,
                "corrected": new_text,
                "context": context_line,
            })

        for opcode, i1, i2, j1, j2 in line_matcher.get_opcodes():
            if opcode == "equal":
                continue

            left_lines = original_lines[i1:i2]
            right_lines = corrected_lines[j1:j2]
            pair_count = max(len(left_lines), len(right_lines))

            for idx in range(pair_count):
                old_line = left_lines[idx] if idx < len(left_lines) else ""
                new_line = right_lines[idx] if idx < len(right_lines) else ""
                corr_line_no = (j1 + idx + 1) if (j1 + idx) < len(corrected_lines) else (i1 + idx + 1)
                is_reference_line = False
                if (j1 + idx) < len(corrected_ref_flags):
                    is_reference_line = corrected_ref_flags[j1 + idx]
                elif (i1 + idx) < len(original_ref_flags):
                    is_reference_line = original_ref_flags[i1 + idx]

                for change in self._line_level_changes(old_line, new_line, is_reference_line):
                    push_item(
                        kind=change["type"],
                        line_number=corr_line_no,
                        old_text=change["original"],
                        new_text=change["corrected"],
                        context_line=new_line or old_line,
                    )

        counts = {kind: len(items) for kind, items in groups.items()}
        total = sum(counts.values())
        return {
            "total": total,
            "counts": counts,
            "groups": groups,
        }

    def _normalize_group_decisions(self, group_decisions: Optional[Dict]) -> Dict[str, bool]:
        """Normalize correction-group accept/reject map to stable booleans."""
        groups = ("spelling", "capitalization", "punctuation", "citation", "reference", "style")
        normalized: Dict[str, bool] = {key: True for key in groups}
        if not isinstance(group_decisions, dict):
            return normalized

        for key in groups:
            if key not in group_decisions:
                continue
            value = group_decisions.get(key)
            if isinstance(value, bool):
                normalized[key] = value
                continue
            if isinstance(value, (int, float)):
                normalized[key] = bool(value)
                continue
            text = str(value).strip().lower()
            if text in ("1", "true", "yes", "accept", "accepted", "on"):
                normalized[key] = True
            elif text in ("0", "false", "no", "reject", "rejected", "off"):
                normalized[key] = False

        return normalized

    def apply_group_decisions(self, original: str, corrected: str, group_decisions: Optional[Dict]) -> str:
        """Apply correction-group accept/reject decisions and build deterministic accepted text."""
        if not (corrected or "").strip():
            return corrected or ""
        if not (original or "").strip():
            return corrected or ""

        decisions = self._normalize_group_decisions(group_decisions)
        original_lines = (original or "").split('\n')
        corrected_lines = (corrected or "").split('\n')
        original_ref_flags = self._reference_line_flags(original_lines)
        corrected_ref_flags = self._reference_line_flags(corrected_lines)
        matcher = difflib.SequenceMatcher(a=original_lines, b=corrected_lines, autojunk=False)
        merged_lines: List[str] = []

        for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
            if opcode == "equal":
                merged_lines.extend(original_lines[i1:i2])
                continue

            left_lines = original_lines[i1:i2]
            right_lines = corrected_lines[j1:j2]
            pair_count = max(len(left_lines), len(right_lines))

            for idx in range(pair_count):
                old_line = left_lines[idx] if idx < len(left_lines) else ""
                new_line = right_lines[idx] if idx < len(right_lines) else ""

                is_reference_line = False
                if (j1 + idx) < len(corrected_ref_flags):
                    is_reference_line = corrected_ref_flags[j1 + idx]
                elif (i1 + idx) < len(original_ref_flags):
                    is_reference_line = original_ref_flags[i1 + idx]

                line_changes = self._line_level_changes(old_line, new_line, is_reference_line)
                change_type = str(line_changes[0].get("type", "style")) if line_changes else "style"
                accept_change = bool(decisions.get(change_type, True))
                chosen_line = new_line if accept_change else old_line

                if chosen_line != "":
                    merged_lines.append(chosen_line)

        return '\n'.join(merged_lines)

    def _reference_line_flags(self, lines: List[str]) -> List[bool]:
        """Mark which lines are inside references section."""
        flags: List[bool] = []
        in_references = False
        heading_re = re.compile(r'^\s*references?\s*:?\s*$', flags=re.IGNORECASE)
        section_break_re = re.compile(
            r'^\s*(?:appendix|acknowledg(?:e)?ments?|funding|conflicts?\s+of\s+interest|author\s+contributions?)\s*:?\s*$',
            flags=re.IGNORECASE
        )

        for line in lines:
            stripped = line.strip()
            if heading_re.match(stripped):
                in_references = True
                flags.append(True)
                continue
            if in_references and section_break_re.match(stripped):
                in_references = False
                flags.append(False)
                continue
            flags.append(in_references)
        return flags

    def _line_level_changes(self, old_line: str, new_line: str, is_reference_line: bool) -> List[Dict]:
        """Return readable line-level changes for one line pair."""
        if old_line == new_line:
            return []

        old_text = old_line.strip()
        new_text = new_line.strip()
        return [{
            "type": self._classify_change(old_text, new_text, is_reference_line),
            "original": old_text,
            "corrected": new_text,
        }]

    def _classify_change(self, old_text: str, new_text: str, is_reference_line: bool) -> str:
        """Classify a single change into UI-friendly groups."""
        old_clean = (old_text or "").strip()
        new_clean = (new_text or "").strip()
        merged = f"{old_clean} {new_clean}".strip()

        if is_reference_line:
            return "reference"

        if re.search(r'(?i)\b(?:https?|ftp)://|www\.|doi:|10\.\d{4,9}/|\[\s*\d+(?:\s*,\s*\d+)*\s*\]', merged):
            return "citation"

        if old_clean and new_clean and old_clean.lower() == new_clean.lower() and old_clean != new_clean:
            return "capitalization"

        old_alpha = re.sub(r'[^A-Za-z]+', '', old_clean).lower()
        new_alpha = re.sub(r'[^A-Za-z]+', '', new_clean).lower()
        if old_alpha and new_alpha and old_alpha == new_alpha and old_clean != new_clean:
            return "punctuation"

        nonword_only = re.sub(r'[\w\s]', '', f"{old_clean}{new_clean}", flags=re.UNICODE)
        has_word_chars = bool(re.search(r'[\w]', f"{old_clean}{new_clean}", flags=re.UNICODE))
        if nonword_only and not has_word_chars:
            return "punctuation"

        if re.fullmatch(r"[A-Za-z][A-Za-z'’\-]*", old_clean) and re.fullmatch(r"[A-Za-z][A-Za-z'’\-]*", new_clean):
            return "spelling"

        return "style"

    def _tokenize_for_diff(self, text: str) -> List[str]:
        """Split text into fine-grained tokens (whitespace, words, punctuation)."""
        # Keeps whitespace as tokens so output spacing/newlines are preserved,
        # while separating punctuation to avoid whole-line highlights.
        return re.findall(r'\s+|[\w⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻]+|[^\w\s]', text, flags=re.UNICODE)

    def _iter_refined_replace_segments(self, deleted_text: str, inserted_text: str):
        """Refine replace chunks to avoid over-highlighting large spans."""
        deleted_tokens = self._tokenize_for_diff(deleted_text)
        inserted_tokens = self._tokenize_for_diff(inserted_text)
        matcher = difflib.SequenceMatcher(a=deleted_tokens, b=inserted_tokens, autojunk=False)

        for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
            if opcode == "equal":
                equal_text = "".join(deleted_tokens[i1:i2])
                if equal_text:
                    yield "equal", equal_text
            elif opcode == "delete":
                del_text = "".join(deleted_tokens[i1:i2])
                if del_text:
                    yield "delete", del_text
            elif opcode == "insert":
                ins_text = "".join(inserted_tokens[j1:j2])
                if ins_text:
                    yield "insert", ins_text
            elif opcode == "replace":
                del_text = "".join(deleted_tokens[i1:i2])
                ins_text = "".join(inserted_tokens[j1:j2])
                if del_text:
                    yield "delete", del_text
                if ins_text:
                    yield "insert", ins_text

    def _iter_diff_segments(self, original: str, corrected: str):
        """Yield (segment_type, segment_text) where type is equal/delete/insert."""
        original_tokens = self._tokenize_for_diff(original)
        corrected_tokens = self._tokenize_for_diff(corrected)
        matcher = difflib.SequenceMatcher(a=original_tokens, b=corrected_tokens, autojunk=False)

        for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
            if opcode == "equal":
                yield "equal", "".join(original_tokens[i1:i2])
            elif opcode == "delete":
                yield "delete", "".join(original_tokens[i1:i2])
            elif opcode == "insert":
                yield "insert", "".join(corrected_tokens[j1:j2])
            elif opcode == "replace":
                deleted = "".join(original_tokens[i1:i2])
                inserted = "".join(corrected_tokens[j1:j2])
                if deleted and inserted:
                    # Refine large replace spans so only changed words are highlighted.
                    yield from self._iter_refined_replace_segments(deleted, inserted)
                else:
                    if deleted:
                        yield "delete", deleted
                    if inserted:
                        yield "insert", inserted

    def _compute_diff(self, original: str, corrected: str) -> List[Dict]:
        """Compute differences between original and corrected."""
        corrections = []
        orig_words = original.split()
        corr_words = corrected.split()

        for i, (orig, corr) in enumerate(zip(orig_words, corr_words)):
            if orig != corr:
                corrections.append({
                    'position': i,
                    'original': orig,
                    'corrected': corr,
                    'type': 'replacement'
                })

        return corrections

    def _create_run_with_format(self, doc, para, text, format_dict):
        """Create a run with specific formatting."""
        run = para.add_run(text + ' ')
        if format_dict.get('strikethrough'):
            run.font.strike = True
        if format_dict.get('underline'):
            run.font.underline = True
        if format_dict.get('color'):
            run.font.color.rgb = format_dict['color']
        if format_dict.get('highlight'):
            run.font.highlight_color = format_dict['highlight']
        return run

    def _add_formatted_run(self, doc, run, text, **kwargs):
        """Add a formatted run to a paragraph."""
        run = doc.add_paragraph().add_run(text)
        if kwargs.get('strikethrough'):
            run.font.strike = True
        if kwargs.get('underline'):
            run.font.underline = True
        if kwargs.get('color'):
            run.font.color.rgb = kwargs['color']
        return run

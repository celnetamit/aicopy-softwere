"""Tests for structure-preserving DOCX load/export behavior."""

import base64
import os
import tempfile
import unittest
import zipfile

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches

from document_processor import DocumentProcessor


class DocxStructurePreservationTests(unittest.TestCase):
    PNG_1X1 = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn8n0sAAAAASUVORK5CYII="
    )

    def setUp(self):
        self.processor = DocumentProcessor()

    def _build_source_docx(self, path: str):
        doc = Document()
        doc.add_paragraph("Intro paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        doc.add_paragraph("Closing paragraph")
        doc.save(path)

    def _build_source_docx_with_image(self, path: str):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as image_handle:
            image_handle.write(self.PNG_1X1)
            image_path = image_handle.name

        try:
            doc = Document()
            doc.add_paragraph("Intro before figure")
            doc.add_picture(image_path, width=Inches(1.0))
            doc.add_paragraph("Figure 1. Original caption")
            doc.add_paragraph("Closing after figure")
            doc.save(path)
        finally:
            os.unlink(image_path)

    def _build_source_docx_with_multi_paragraph_cell(self, path: str):
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=1, cols=2)
        left = table.cell(0, 0)
        left.paragraphs[0].text = "Cell A line 1"
        left.add_paragraph("Cell A line 2")
        right = table.cell(0, 1)
        right.paragraphs[0].text = "Cell B"
        doc.add_paragraph("After table")
        doc.save(path)

    def _build_source_docx_with_nested_table_cell(self, path: str):
        doc = Document()
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Header text"
        section.footer.paragraphs[0].text = "Footer text"
        doc.add_paragraph("Before nested table")
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = "Outer intro"
        nested = cell.add_table(rows=1, cols=2)
        nested.cell(0, 0).text = "Inner A"
        nested.cell(0, 1).text = "Inner B"
        cell.add_paragraph("Outer outro")
        doc.add_paragraph("After nested table")
        doc.save(path)

    def _build_source_docx_with_header_footer(self, path: str):
        doc = Document()
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Running header"
        section.footer.paragraphs[0].text = "Page footer"
        doc.add_paragraph("Body paragraph one")
        doc.add_paragraph("Body paragraph two")
        doc.save(path)

    def _append_vml_textbox(self, paragraph, textbox_text: str):
        textbox_xml = (
            f'<w:r {nsdecls("w")} xmlns:v="urn:schemas-microsoft-com:vml">'
            '<w:pict>'
            '<v:shape id="TextBox1" style="width:200pt;height:40pt" type="#_x0000_t202">'
            '<v:textbox>'
            '<w:txbxContent>'
            '<w:p><w:r><w:t xml:space="preserve">'
            f'{textbox_text}'
            '</w:t></w:r></w:p>'
            '</w:txbxContent>'
            '</v:textbox>'
            '</v:shape>'
            '</w:pict>'
            '</w:r>'
        )
        paragraph._p.append(parse_xml(textbox_xml))

    def _build_source_docx_with_textbox(self, path: str):
        doc = Document()
        doc.add_paragraph("Before textbox")
        paragraph = doc.add_paragraph("Body beside box")
        self._append_vml_textbox(paragraph, "Textbox original")
        doc.add_paragraph("After textbox")
        doc.save(path)

    def _augment_docx_with_reference_parts(self, path: str):
        with zipfile.ZipFile(path, "a") as archive:
            archive.writestr(
                "word/comments.xml",
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:comment w:id="0" w:author="Tester" w:date="2026-04-18T00:00:00Z"><w:p><w:r><w:t>Comment text</w:t></w:r></w:p></w:comment>'
                "</w:comments>",
            )
            archive.writestr(
                "word/footnotes.xml",
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:id="-1"/><w:footnote w:id="0"/><w:footnote w:id="1"><w:p><w:r><w:t>Footnote text</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>",
            )
            archive.writestr(
                "word/endnotes.xml",
                '<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:endnote w:id="-1"/><w:endnote w:id="0"/><w:endnote w:id="1"><w:p><w:r><w:t>Endnote text</w:t></w:r></w:p></w:endnote>'
                "</w:endnotes>",
            )

    def _textbox_texts(self, paragraph) -> list[str]:
        values = []
        try:
            containers = paragraph._p.xpath('.//*[local-name()="txbxContent"]')
        except Exception:
            containers = []
        for container in containers:
            texts = container.xpath('.//*[local-name()="t"]/text()')
            values.append("".join(texts))
        return values

    def test_load_document_includes_table_rows(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            source_path = handle.name

        try:
            self._build_source_docx(source_path)
            text, fmt = self.processor.load_document(source_path)
            self.assertEqual(fmt, "docx")
            self.assertIn("Intro paragraph", text)
            self.assertIn("A1\tB1", text)
            self.assertIn("A2\tB2", text)
            self.assertIn("Closing paragraph", text)
        finally:
            os.unlink(source_path)

    def test_generate_clean_docx_preserves_table_structure_from_source_docx(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx(source_path)
            corrected = "\n".join([
                "Updated intro paragraph",
                "R1C1\tR1C2",
                "R2C1\tR2C2",
                "Updated closing paragraph",
            ])

            self.processor.generate_clean_docx(corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            self.assertEqual(len(out_doc.tables), 1)
            self.assertEqual(out_doc.tables[0].cell(0, 0).text, "R1C1")
            self.assertEqual(out_doc.tables[0].cell(0, 1).text, "R1C2")
            self.assertEqual(out_doc.tables[0].cell(1, 0).text, "R2C1")
            self.assertEqual(out_doc.tables[0].cell(1, 1).text, "R2C2")
            self.assertEqual(out_doc.paragraphs[0].text, "Updated intro paragraph")
            self.assertEqual(out_doc.paragraphs[-1].text, "Updated closing paragraph")
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_load_document_skips_image_only_paragraphs_but_keeps_caption_order(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            source_path = handle.name

        try:
            self._build_source_docx_with_image(source_path)
            text, fmt = self.processor.load_document(source_path)
            self.assertEqual(fmt, "docx")
            lines = text.split("\n")
            self.assertEqual(lines, [
                "Intro before figure",
                "Figure 1. Original caption",
                "Closing after figure",
            ])
        finally:
            os.unlink(source_path)

    def test_generate_clean_docx_preserves_image_and_caption_alignment(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_image(source_path)
            corrected = "\n".join([
                "Updated intro before figure",
                "Figure 1. Updated caption",
                "Updated closing after figure",
            ])

            self.processor.generate_clean_docx(corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            self.assertEqual(len(out_doc.inline_shapes), 1)
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]
            self.assertIn("Updated intro before figure", paragraph_texts)
            self.assertIn("Figure 1. Updated caption", paragraph_texts)
            self.assertIn("Updated closing after figure", paragraph_texts)
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_load_document_preserves_multi_paragraph_table_cell_markers(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            source_path = handle.name

        try:
            self._build_source_docx_with_multi_paragraph_cell(source_path)
            text, fmt = self.processor.load_document(source_path)
            self.assertEqual(fmt, "docx")
            self.assertIn("Cell A line 1[[CELL_PARA]]Cell A line 2\tCell B", text)
        finally:
            os.unlink(source_path)

    def test_generate_clean_docx_preserves_multi_paragraph_table_cells(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_multi_paragraph_cell(source_path)
            corrected = "\n".join([
                "Updated before table",
                "Updated A1[[CELL_PARA]]Updated A2\tUpdated B",
                "Updated after table",
            ])

            self.processor.generate_clean_docx(corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            cell = out_doc.tables[0].cell(0, 0)
            self.assertEqual([p.text for p in cell.paragraphs], ["Updated A1", "Updated A2"])
            self.assertEqual(out_doc.tables[0].cell(0, 1).text, "Updated B")
            self.assertEqual(out_doc.paragraphs[0].text, "Updated before table")
            self.assertEqual(out_doc.paragraphs[-1].text, "Updated after table")
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_load_document_serializes_nested_table_content_inside_cells(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            source_path = handle.name

        try:
            self._build_source_docx_with_nested_table_cell(source_path)
            text, fmt = self.processor.load_document(source_path)
            self.assertEqual(fmt, "docx")
            self.assertIn(
                "P:Outer intro[[CELL_BLOCK]]T:Inner A[[CELL_TABLE_CELL]]Inner B[[CELL_BLOCK]]P:Outer outro",
                text,
            )
        finally:
            os.unlink(source_path)

    def test_generate_clean_docx_preserves_nested_table_cells(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_nested_table_cell(source_path)
            corrected = "\n".join([
                "Updated before nested table",
                "P:Updated outer intro[[CELL_BLOCK]]T:Updated inner A[[CELL_TABLE_CELL]]Updated inner B[[CELL_BLOCK]]P:Updated outer outro",
                "Updated after nested table",
            ])

            self.processor.generate_clean_docx(corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            outer_cell = out_doc.tables[0].cell(0, 0)
            self.assertEqual(outer_cell.paragraphs[0].text, "Updated outer intro")
            self.assertEqual(outer_cell.paragraphs[-1].text, "Updated outer outro")
            self.assertEqual(len(outer_cell.tables), 1)
            self.assertEqual(outer_cell.tables[0].cell(0, 0).text, "Updated inner A")
            self.assertEqual(outer_cell.tables[0].cell(0, 1).text, "Updated inner B")
            self.assertEqual(out_doc.paragraphs[0].text, "Updated before nested table")
            self.assertEqual(out_doc.paragraphs[-1].text, "Updated after nested table")
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_generate_highlighted_docx_preserves_image_and_caption_alignment(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_image(source_path)
            corrected = "\n".join([
                "Updated intro before figure",
                "Figure 1. Updated caption",
                "Updated closing after figure",
            ])

            self.processor.generate_highlighted_docx("", corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            self.assertEqual(len(out_doc.inline_shapes), 1)
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]
            self.assertTrue(any("Intro" in text and "Updated intro before figure" in text for text in paragraph_texts))
            self.assertTrue(any("Figure 1." in text and "Updated caption" in text for text in paragraph_texts))
            self.assertTrue(any("Closing" in text and "Updated closing after figure" in text for text in paragraph_texts))
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_generate_highlighted_docx_preserves_nested_table_cells(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_nested_table_cell(source_path)
            corrected = "\n".join([
                "Updated before nested table",
                "P:Updated outer intro[[CELL_BLOCK]]T:Updated inner A[[CELL_TABLE_CELL]]Updated inner B[[CELL_BLOCK]]P:Updated outer outro",
                "Updated after nested table",
            ])

            self.processor.generate_highlighted_docx("", corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            outer_cell = out_doc.tables[0].cell(0, 0)
            self.assertEqual(len(outer_cell.tables), 1)
            self.assertIn("Outer", outer_cell.paragraphs[0].text)
            self.assertIn("Updated outer intro", outer_cell.paragraphs[0].text)
            self.assertIn("Outer", outer_cell.paragraphs[-1].text)
            self.assertIn("Updated outer outro", outer_cell.paragraphs[-1].text)
            self.assertIn("Inner", outer_cell.tables[0].cell(0, 0).text)
            self.assertIn("Updated inner A", outer_cell.tables[0].cell(0, 0).text)
            self.assertIn("Inner", outer_cell.tables[0].cell(0, 1).text)
            self.assertIn("Updated inner B", outer_cell.tables[0].cell(0, 1).text)
            self.assertIn("Header text", out_doc.sections[0].header.paragraphs[0].text)
            self.assertIn("Footer text", out_doc.sections[0].footer.paragraphs[0].text)
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_generate_clean_docx_preserves_header_and_footer(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_header_footer(source_path)
            corrected = "\n".join([
                "Updated body paragraph one",
                "Updated body paragraph two",
            ])

            self.processor.generate_clean_docx(corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            self.assertEqual(out_doc.sections[0].header.paragraphs[0].text, "Running header")
            self.assertEqual(out_doc.sections[0].footer.paragraphs[0].text, "Page footer")
            self.assertEqual(out_doc.paragraphs[0].text, "Updated body paragraph one")
            self.assertEqual(out_doc.paragraphs[1].text, "Updated body paragraph two")
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_generate_highlighted_docx_preserves_header_and_footer(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_header_footer(source_path)
            corrected = "\n".join([
                "Updated body paragraph one",
                "Updated body paragraph two",
            ])

            self.processor.generate_highlighted_docx("", corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            self.assertEqual(out_doc.sections[0].header.paragraphs[0].text, "Running header")
            self.assertEqual(out_doc.sections[0].footer.paragraphs[0].text, "Page footer")
            self.assertIn("Body", out_doc.paragraphs[0].text)
            self.assertIn("Updated body paragraph one", out_doc.paragraphs[0].text)
            self.assertIn("Body", out_doc.paragraphs[1].text)
            self.assertIn("Updated body paragraph two", out_doc.paragraphs[1].text)
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_load_document_includes_textbox_content(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            source_path = handle.name

        try:
            self._build_source_docx_with_textbox(source_path)
            text, fmt = self.processor.load_document(source_path)
            self.assertEqual(fmt, "docx")
            self.assertIn("Body beside box[[TEXTBOX]]Textbox original", text)
        finally:
            os.unlink(source_path)

    def test_generate_clean_docx_updates_textbox_content(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_textbox(source_path)
            corrected = "\n".join([
                "Before textbox updated",
                "Body beside box updated[[TEXTBOX]]Textbox updated",
                "After textbox updated",
            ])

            self.processor.generate_clean_docx(corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            self.assertEqual(out_doc.paragraphs[0].text, "Before textbox updated")
            self.assertEqual(out_doc.paragraphs[1].text, "Body beside box updated")
            self.assertEqual(self._textbox_texts(out_doc.paragraphs[1]), ["Textbox updated"])
            self.assertEqual(out_doc.paragraphs[2].text, "After textbox updated")
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_generate_highlighted_docx_updates_textbox_content(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx_with_textbox(source_path)
            corrected = "\n".join([
                "Before textbox updated",
                "Body beside box updated[[TEXTBOX]]Textbox updated",
                "After textbox updated",
            ])

            self.processor.generate_highlighted_docx("", corrected, output_path, source_docx_path=source_path)

            out_doc = Document(output_path)
            self.assertIn("Body", out_doc.paragraphs[1].text)
            self.assertIn("updated", out_doc.paragraphs[1].text)
            textbox_values = self._textbox_texts(out_doc.paragraphs[1])
            self.assertEqual(len(textbox_values), 1)
            self.assertIn("Textbox original", textbox_values[0])
            self.assertIn("updated", textbox_values[0])
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_process_text_attaches_docx_package_feature_summary(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            source_path = handle.name

        try:
            self._build_source_docx_with_textbox(source_path)
            self._augment_docx_with_reference_parts(source_path)
            text, fmt = self.processor.load_document(source_path)
            self.assertEqual(fmt, "docx")
            self.processor.process_text(
                text,
                {
                    "spelling": True,
                    "sentence_case": True,
                    "punctuation": True,
                    "chicago_style": True,
                    "ai": {"enabled": False},
                },
            )
            audit = self.processor.get_processing_audit()
            package = audit.get("summary", {}).get("docx_package_features", {})
            self.assertEqual(package.get("comments"), 1)
            self.assertEqual(package.get("footnotes"), 1)
            self.assertEqual(package.get("endnotes"), 1)
            self.assertEqual(package.get("textboxes"), 1)
        finally:
            os.unlink(source_path)

    def test_generate_clean_docx_preserves_reference_package_parts(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            self._build_source_docx(source_path)
            self._augment_docx_with_reference_parts(source_path)

            self.processor.generate_clean_docx(
                "Updated intro paragraph\nR1C1\tR1C2\nR2C1\tR2C2\nUpdated closing paragraph",
                output_path,
                source_docx_path=source_path,
            )

            with zipfile.ZipFile(output_path) as archive:
                names = set(archive.namelist())
                self.assertIn("word/comments.xml", names)
                self.assertIn("word/footnotes.xml", names)
                self.assertIn("word/endnotes.xml", names)
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_generate_clean_docx_fallback_preserves_blank_lines_and_structure_styles(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            corrected = "\n".join([
                "# Executive Summary",
                "",
                "- First bullet",
                "1. First numbered item",
                "Body paragraph.",
            ])

            self.processor.generate_clean_docx(corrected, output_path)

            out_doc = Document(output_path)
            self.assertEqual(len(out_doc.paragraphs), 5)
            self.assertEqual(out_doc.paragraphs[0].text, "Executive Summary")
            self.assertEqual(out_doc.paragraphs[0].style.name, "Heading 1")
            self.assertEqual(out_doc.paragraphs[1].text, "")
            self.assertEqual(out_doc.paragraphs[2].text, "First bullet")
            self.assertEqual(out_doc.paragraphs[2].style.name, "List Bullet")
            self.assertEqual(out_doc.paragraphs[3].text, "First numbered item")
            self.assertEqual(out_doc.paragraphs[3].style.name, "List Number")
            self.assertEqual(out_doc.paragraphs[4].text, "Body paragraph.")
        finally:
            os.unlink(output_path)

    def test_generate_highlighted_docx_fallback_uses_structured_paragraph_styles(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        try:
            original = "\n".join([
                "# Executive Summary",
                "",
                "- Old bullet",
                "Body paragraph.",
            ])
            corrected = "\n".join([
                "# Executive Summary",
                "",
                "- New bullet",
                "Body paragraph updated.",
            ])

            self.processor.generate_highlighted_docx(original, corrected, output_path)

            out_doc = Document(output_path)
            self.assertEqual(out_doc.paragraphs[0].style.name, "Heading 1")
            self.assertEqual(out_doc.paragraphs[0].text, "Executive Summary")
            self.assertEqual(out_doc.paragraphs[1].text, "")
            self.assertEqual(out_doc.paragraphs[2].style.name, "List Bullet")
            self.assertIn("Old", out_doc.paragraphs[2].text)
            self.assertIn("New bullet", out_doc.paragraphs[2].text)
            self.assertIn("Body paragraph", out_doc.paragraphs[3].text)
            self.assertIn("updated", out_doc.paragraphs[3].text)
        finally:
            os.unlink(output_path)


if __name__ == "__main__":
    unittest.main()

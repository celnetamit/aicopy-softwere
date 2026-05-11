"""WSGI tests for authenticated Manuscript Editor web app."""

import base64
import io
import json
import math
import os
import tempfile
import unittest
import zipfile
from unittest.mock import Mock, patch
from urllib.parse import urlencode
from wsgiref.util import setup_testing_defaults

from docx import Document

from chicago_editor import ChicagoEditor


os.environ.setdefault("MANUSCRIPT_EDITOR_DEV_TEST_TOKENS", "1")
os.environ.setdefault("DATABASE_URL", "sqlite:////tmp/manuscript_editor_test.sqlite3")
os.environ.setdefault("GOOGLE_CLIENT_ID", "test-google-client-id")

import webapp  # noqa: E402


class WsgiTestClient:
    def __init__(self, app):
        self.app = app
        self.cookies = {}

    def request(self, method, path, payload=None, query=None, headers=None):
        body = b""
        if payload is not None:
            body = json.dumps(payload).encode("utf-8")

        environ = {}
        setup_testing_defaults(environ)
        environ["REQUEST_METHOD"] = method.upper()

        path_info = path
        query_string = ""
        if "?" in path:
            path_info, query_string = path.split("?", 1)
        if query:
            encoded = urlencode(query)
            query_string = f"{query_string}&{encoded}" if query_string else encoded

        environ["PATH_INFO"] = path_info
        environ["QUERY_STRING"] = query_string
        environ["CONTENT_LENGTH"] = str(len(body))
        environ["wsgi.input"] = io.BytesIO(body)

        if body:
            environ["CONTENT_TYPE"] = "application/json"

        if self.cookies:
            environ["HTTP_COOKIE"] = "; ".join(f"{key}={value}" for key, value in self.cookies.items())

        if headers:
            for raw_name, raw_value in headers.items():
                if raw_name is None or raw_value is None:
                    continue
                name = str(raw_name).strip()
                if not name:
                    continue
                key = name.upper().replace("-", "_")
                if key not in ("CONTENT_TYPE", "CONTENT_LENGTH") and not key.startswith("HTTP_"):
                    key = "HTTP_" + key
                environ[key] = str(raw_value)

        meta = {}

        def start_response(status, headers, exc_info=None):
            meta["status"] = status
            meta["headers"] = headers

        result = self.app(environ, start_response)
        response_body = b"".join(result)
        if hasattr(result, "close"):
            result.close()

        for header_name, header_value in meta.get("headers", []):
            if header_name.lower() != "set-cookie":
                continue
            cookie_pair = header_value.split(";", 1)[0]
            cookie_name, cookie_value = cookie_pair.split("=", 1)
            self.cookies[cookie_name] = cookie_value

        text = response_body.decode("utf-8") if response_body else ""
        data = json.loads(text) if text else {}
        status_code = int(str(meta.get("status", "500")).split(" ", 1)[0])
        return status_code, data

    def request_text(self, method, path, payload=None, query=None, headers=None):
        body = b""
        if payload is not None:
            body = json.dumps(payload).encode("utf-8")

        environ = {}
        setup_testing_defaults(environ)
        environ["REQUEST_METHOD"] = method.upper()

        path_info = path
        query_string = ""
        if "?" in path:
            path_info, query_string = path.split("?", 1)
        if query:
            encoded = urlencode(query)
            query_string = f"{query_string}&{encoded}" if query_string else encoded

        environ["PATH_INFO"] = path_info
        environ["QUERY_STRING"] = query_string
        environ["CONTENT_LENGTH"] = str(len(body))
        environ["wsgi.input"] = io.BytesIO(body)

        if body:
            environ["CONTENT_TYPE"] = "application/json"

        if self.cookies:
            environ["HTTP_COOKIE"] = "; ".join(f"{key}={value}" for key, value in self.cookies.items())

        if headers:
            for raw_name, raw_value in headers.items():
                if raw_name is None or raw_value is None:
                    continue
                name = str(raw_name).strip()
                if not name:
                    continue
                key = name.upper().replace("-", "_")
                if key not in ("CONTENT_TYPE", "CONTENT_LENGTH") and not key.startswith("HTTP_"):
                    key = "HTTP_" + key
                environ[key] = str(raw_value)

        meta = {}

        def start_response(status, headers, exc_info=None):
            meta["status"] = status
            meta["headers"] = headers

        result = self.app(environ, start_response)
        response_body = b"".join(result)
        if hasattr(result, "close"):
            result.close()

        status_code = int(str(meta.get("status", "500")).split(" ", 1)[0])
        text = response_body.decode("utf-8") if response_body else ""
        return status_code, text

    def request_bytes(self, method, path, payload=None, query=None, headers=None):
        body = b""
        if payload is not None:
            body = json.dumps(payload).encode("utf-8")

        environ = {}
        setup_testing_defaults(environ)
        environ["REQUEST_METHOD"] = method.upper()

        path_info = path
        query_string = ""
        if "?" in path:
            path_info, query_string = path.split("?", 1)
        if query:
            encoded = urlencode(query)
            query_string = f"{query_string}&{encoded}" if query_string else encoded

        environ["PATH_INFO"] = path_info
        environ["QUERY_STRING"] = query_string
        environ["CONTENT_LENGTH"] = str(len(body))
        environ["wsgi.input"] = io.BytesIO(body)

        if body:
            environ["CONTENT_TYPE"] = "application/json"

        if self.cookies:
            environ["HTTP_COOKIE"] = "; ".join(f"{key}={value}" for key, value in self.cookies.items())

        if headers:
            for raw_name, raw_value in headers.items():
                if raw_name is None or raw_value is None:
                    continue
                name = str(raw_name).strip()
                if not name:
                    continue
                key = name.upper().replace("-", "_")
                if key not in ("CONTENT_TYPE", "CONTENT_LENGTH") and not key.startswith("HTTP_"):
                    key = "HTTP_" + key
                environ[key] = str(raw_value)

        meta = {}

        def start_response(status, headers, exc_info=None):
            meta["status"] = status
            meta["headers"] = headers

        result = self.app(environ, start_response)
        response_body = b"".join(result)
        if hasattr(result, "close"):
            result.close()

        status_code = int(str(meta.get("status", "500")).split(" ", 1)[0])
        headers_dict = {str(k).lower(): str(v) for k, v in meta.get("headers", [])}
        return status_code, response_body, headers_dict


class AuthenticatedWebAppApiTests(unittest.TestCase):
    def setUp(self):
        webapp._STORE.clear_all_for_tests()
        with ChicagoEditor._SHARED_ONLINE_VALIDATION_CACHE_LOCK:
            ChicagoEditor._SHARED_ONLINE_VALIDATION_CACHE.clear()
        with ChicagoEditor._SHARED_ONLINE_LOOKUP_METRICS_LOCK:
            ChicagoEditor._SHARED_ONLINE_LOOKUP_METRICS = ChicagoEditor._default_online_lookup_metrics()
            ChicagoEditor._SHARED_ONLINE_LOOKUP_METRICS_UPDATED_AT = 0
        self.client = WsgiTestClient(webapp.app)
        self._old_local_login_enabled = webapp.ENABLE_LOCAL_MANUAL_LOGIN

    def tearDown(self):
        webapp.ENABLE_LOCAL_MANUAL_LOGIN = self._old_local_login_enabled

    def _login(self, email="user@conwiz.in"):
        status, payload = self.client.request(
            "POST",
            "/api/auth/google-login",
            {"id_token": f"test:{email}"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertIn("user", payload)
        return payload["user"]

    def test_blocked_domain_cannot_login(self):
        status, payload = self.client.request(
            "POST",
            "/api/auth/google-login",
            {"id_token": "test:blocked@example.com"},
        )
        self.assertEqual(status, 403)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "AUTH_DOMAIN_BLOCKED")

    def test_auth_me_requires_login_then_succeeds(self):
        status, payload = self.client.request("GET", "/api/auth/me")
        self.assertEqual(status, 401)
        self.assertFalse(payload.get("success"))

        user = self._login("staff@conwiz.in")
        status, payload = self.client.request("GET", "/api/auth/me")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload["user"]["email"], "staff@conwiz.in")
        self.assertEqual(payload["user"]["role"], user["role"])

    def test_local_manual_login_disabled_by_default(self):
        status, payload = self.client.request(
            "POST",
            "/api/auth/local-login",
            {"username": "admin", "password": "password"},
        )
        self.assertEqual(status, 403)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "AUTH_LOCAL_LOGIN_DISABLED")

    def test_local_manual_login_succeeds_when_enabled(self):
        webapp.ENABLE_LOCAL_MANUAL_LOGIN = True
        status, payload = self.client.request(
            "POST",
            "/api/auth/local-login",
            {"username": "admin", "password": "password"},
            headers={"X-Forwarded-For": "127.0.0.1"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload["user"]["role"], "ADMIN")

        status, payload = self.client.request("GET", "/api/auth/me")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload["user"]["role"], "ADMIN")

    def test_admin_dashboard_route_renders_admin_shell(self):
        status, html = self.client.request_text("GET", "/admin-dashboard")
        self.assertEqual(status, 200)
        self.assertIn('id="admin-panel-backdrop"', html)
        self.assertIn('admin-dashboard-active', html)
        self.assertIn('class="setup-wizard-backdrop" id="admin-panel-backdrop"', html)
        self.assertNotIn('class="setup-wizard-backdrop hidden" id="admin-panel-backdrop"', html)
        self.assertNotIn("{{", html)

    def test_root_redirects_to_tasks_dashboard(self):
        environ = {}
        setup_testing_defaults(environ)
        environ["REQUEST_METHOD"] = "GET"
        environ["PATH_INFO"] = "/"
        meta = {}

        def start_response(status, headers, exc_info=None):
            meta["status"] = status
            meta["headers"] = headers

        result = webapp.app(environ, start_response)
        body = b"".join(result)
        if hasattr(result, "close"):
            result.close()

        self.assertEqual(int(str(meta.get("status", "500")).split(" ", 1)[0]), 302)
        headers = dict(meta.get("headers", []))
        self.assertEqual(headers.get("Location"), "/tasks")
        self.assertEqual(body, b"")

    def test_tasks_dashboard_route_renders_dashboard_shell(self):
        status, html = self.client.request_text("GET", "/tasks")
        self.assertEqual(status, 200)
        self.assertIn('class="tasks-dashboard-route"', html)
        self.assertIn("<title>Manuscript Editor - Tasks</title>", html)
        self.assertIn("Task Dashboard", html)
        self.assertNotIn('id="preview-text"', html)
        self.assertNotIn('id="process-btn"', html)
        self.assertNotIn("{{", html)

    def test_task_detail_route_renders_task_detail_shell(self):
        status, html = self.client.request_text("GET", "/tasks/example-task-id")
        self.assertEqual(status, 200)
        self.assertIn('class="task-detail-route"', html)
        self.assertIn('data-task-route-id="example-task-id"', html)
        self.assertIn("<title>Manuscript Editor - Task Detail</title>", html)
        self.assertIn('id="preview-text"', html)
        self.assertIn('id="process-btn"', html)
        self.assertNotIn("{{", html)

    def test_task_detail_rerun_unresolved_button_has_safe_label_and_tooltip(self):
        status, html = self.client.request_text("GET", "/tasks/example-task-id")
        self.assertEqual(status, 200)
        self.assertIn('id="rerun-unresolved-btn"', html)
        self.assertIn("Rerun Unresolved Refs (Safe)", html)
        self.assertIn("Rerun only unresolved references using safe settings", html)
        self.assertIn('id="assistant-unresolved-panel"', html)
        self.assertIn('id="assistant-unresolved-sort"', html)
        self.assertIn('id="assistant-unresolved-rerun-btn"', html)
        self.assertIn('id="assistant-unresolved-rerun-autofixable-btn"', html)
        self.assertIn('id="assistant-export-unresolved-btn"', html)

    def test_admin_dashboard_contains_new_reference_automation_controls(self):
        status, html = self.client.request_text("GET", "/admin-dashboard")
        self.assertEqual(status, 200)
        self.assertIn('id="admin-setting-online-reference-validation-admin-cap"', html)
        self.assertIn('id="admin-setting-auto-resolve-unresolved-references"', html)
        self.assertIn('id="admin-reference-unresolved-trend-summary"', html)
        self.assertIn('id="admin-reference-diagnostics-trends-output"', html)

    def test_version_endpoint_and_footer_use_shared_version_source(self):
        version_file = os.path.join(os.path.dirname(__file__), "..", "VERSION")
        with open(version_file, "r", encoding="utf-8") as handle:
            expected_version = handle.read().strip()
        self.assertTrue(expected_version)

        status, payload = self.client.request("GET", "/api/version")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload.get("version"), expected_version)
        self.assertEqual(payload.get("asset_version"), f"v{expected_version}")

        status, html = self.client.request_text("GET", "/tasks/example-task-id")
        self.assertEqual(status, 200)
        self.assertIn(f"v{expected_version}", html)

    def test_packaging_defaults_read_shared_version_file(self):
        version_file = os.path.join(os.path.dirname(__file__), "..", "VERSION")
        with open(version_file, "r", encoding="utf-8") as handle:
            expected_version = handle.read().strip()
        self.assertTrue(expected_version)

        linux_script_path = os.path.join(os.path.dirname(__file__), "..", "scripts", "linux", "build_deb.sh")
        with open(linux_script_path, "r", encoding="utf-8") as handle:
            linux_script = handle.read()
        self.assertIn('DEFAULT_VERSION="$(tr -d', linux_script)
        self.assertIn('${ROOT_DIR}/VERSION', linux_script)
        self.assertIn('VERSION="${1:-${DEFAULT_VERSION}}"', linux_script)

        windows_iss_path = os.path.join(os.path.dirname(__file__), "..", "packaging", "windows", "ManuscriptEditor.iss")
        with open(windows_iss_path, "r", encoding="utf-8") as handle:
            windows_iss = handle.read()
        self.assertIn('#define MyAppVersion Trim(FileRead("..\\..\\VERSION"))', windows_iss)

    def test_reference_validation_diagnostics_includes_admin_cap_setting(self):
        self._login("admin@conwiz.in")
        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request(
            "POST",
            "/api/auth/google-login",
            {"id_token": "test:amit@conwiz.in"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        status, payload = admin_client.request("GET", "/api/admin/reference-validation-diagnostics")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        diagnostics = payload.get("diagnostics") or {}
        runtime = diagnostics.get("global_runtime") or {}
        self.assertIn("online_reference_validation_admin_cap", runtime)
        self.assertEqual(int(runtime.get("online_reference_validation_admin_cap", 0)), 150)
        self.assertIn("unresolved_trends", diagnostics)

    def test_rerun_unresolved_frontend_has_direct_process_fallback_path(self):
        app_js_path = os.path.join(os.path.dirname(__file__), "..", "web", "app.js")
        with open(app_js_path, "r", encoding="utf-8") as handle:
            source = handle.read()
        self.assertIn("rerun_unresolved_references_fallback", source)
        self.assertIn("assistant_reprocess_task", source)
        self.assertIn("eel.process_document(retryOptions, taskId)", source)
        self.assertIn("Used direct fallback", source)
        self.assertIn("buildUnresolvedRerunDelta", source)
        self.assertIn("Unresolved references rerun complete. Before:", source)

    def test_corrections_panel_renders_rerun_unresolved_delta_chips(self):
        preview_js_path = os.path.join(os.path.dirname(__file__), "..", "web", "app-preview.js")
        with open(preview_js_path, "r", encoding="utf-8") as handle:
            source = handle.read()
        self.assertIn("Unresolved Delta:", source)
        self.assertIn("Rerun Regression:", source)
        self.assertIn("Unresolved reason:", source)

    def test_unresolved_references_panel_actions_are_wired(self):
        app_js_path = os.path.join(os.path.dirname(__file__), "..", "web", "app.js")
        with open(app_js_path, "r", encoding="utf-8") as handle:
            app_source = handle.read()
        self.assertIn("collectUnresolvedReferenceItemsFromState", app_source)
        self.assertIn("exportUnresolvedReferencesReport", app_source)
        self.assertIn("unresolved_references_", app_source)

        settings_js_path = os.path.join(os.path.dirname(__file__), "..", "web", "app-settings.js")
        with open(settings_js_path, "r", encoding="utf-8") as handle:
            settings_source = handle.read()
        self.assertIn("assistantUnresolvedRerunBtn", settings_source)
        self.assertIn("assistantUnresolvedRerunAutofixableBtn", settings_source)
        self.assertIn("assistantExportUnresolvedBtn", settings_source)
        self.assertIn("assistantUnresolvedSort", settings_source)

    def test_json_response_sanitizes_non_finite_numbers(self):
        response = webapp._json_response({"value": math.nan, "nested": {"score": math.inf}})
        payload = json.loads(response.body)
        self.assertIsNone(payload["value"])
        self.assertIsNone(payload["nested"]["score"])

    def test_upload_process_and_download_round_trip(self):
        self._login("writer@conwiz.in")

        status, payload = self.client.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "sample.txt", "content": "This are sample text."},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        task_id = payload.get("task_id")
        self.assertTrue(task_id)

        status, payload = self.client.request(
            "POST",
            f"/api/tasks/{task_id}/process",
            {
                "options": {
                    "spelling": True,
                    "sentence_case": True,
                    "punctuation": True,
                    "chicago_style": True,
                    "ai": {"enabled": False},
                }
            },
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload.get("task_id"), task_id)
        self.assertIn("prose_only_diff", payload)
        self.assertIsInstance(payload.get("prose_only_diff"), str)
        self.assertIn("strict_cmos_issues", payload)
        self.assertIsInstance(payload.get("strict_cmos_issues"), dict)

        status, payload = self.client.request("GET", f"/api/tasks/{task_id}/download", query={"type": "clean"})
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertIn("base64_data", payload)
        exported_bytes = base64.b64decode(payload["base64_data"])
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name
        try:
            with open(output_path, "wb") as outfile:
                outfile.write(exported_bytes)
            self.assertTrue(zipfile.is_zipfile(output_path))
            with zipfile.ZipFile(output_path) as archive:
                names = set(archive.namelist())
                self.assertIn("[Content_Types].xml", names)
                self.assertIn("word/document.xml", names)
                self.assertIn("word/_rels/document.xml.rels", names)
        finally:
            os.unlink(output_path)

        status, payload = self.client.request("GET", "/api/tasks")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertGreaterEqual(len(payload.get("tasks", [])), 1)

    def test_binary_download_endpoint_returns_valid_docx(self):
        self._login("writer@conwiz.in")

        status, payload = self.client.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "sample.txt", "content": "This are sample text."},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        task_id = payload.get("task_id")
        self.assertTrue(task_id)

        status, payload = self.client.request(
            "POST",
            f"/api/tasks/{task_id}/process",
            {
                "options": {
                    "spelling": True,
                    "sentence_case": True,
                    "punctuation": True,
                    "chicago_style": True,
                    "ai": {"enabled": False},
                }
            },
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, body, headers = self.client.request_bytes(
            "GET",
            f"/api/tasks/{task_id}/download-file",
            query={"type": "clean"},
        )
        self.assertEqual(status, 200)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers.get("content-type", ""),
        )
        self.assertIn("attachment;", headers.get("content-disposition", ""))

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name
        try:
            with open(output_path, "wb") as outfile:
                outfile.write(body)
            self.assertTrue(zipfile.is_zipfile(output_path))
            with zipfile.ZipFile(output_path) as archive:
                names = set(archive.namelist())
                self.assertIn("[Content_Types].xml", names)
                self.assertIn("word/document.xml", names)
                self.assertIn("word/_rels/document.xml.rels", names)
        finally:
            os.unlink(output_path)

    def test_legacy_export_file_uses_docx_template_when_base64_source_is_supplied(self):
        self._login("writer@conwiz.in")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name

        try:
            doc = Document()
            doc.add_paragraph("Intro paragraph")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A1"
            table.cell(0, 1).text = "B1"
            doc.add_paragraph("Closing paragraph")
            doc.save(source_path)

            with open(source_path, "rb") as infile:
                source_docx_base64 = base64.b64encode(infile.read()).decode("ascii")

            status, payload = self.client.request(
                "POST",
                "/api/export-file",
                {
                    "task_id": "",
                    "source_type": "docx",
                    "source_docx_base64": source_docx_base64,
                    "file_type": "clean",
                    "original_text": "Intro paragraph\nA1\tB1\nClosing paragraph",
                    "corrected_text": "Updated intro\nR1C1\tR1C2\nUpdated closing",
                    "file_name": "sample.docx",
                },
            )
            self.assertEqual(status, 200)
            self.assertTrue(payload.get("success"))
            self.assertIn("base64_data", payload)

            exported_bytes = base64.b64decode(payload["base64_data"])
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
                output_path = output_handle.name

            try:
                with open(output_path, "wb") as outfile:
                    outfile.write(exported_bytes)
                out_doc = Document(output_path)
                self.assertEqual(out_doc.paragraphs[0].text, "Updated intro")
                self.assertEqual(len(out_doc.tables), 1)
                self.assertEqual(out_doc.tables[0].cell(0, 0).text, "R1C1")
                self.assertEqual(out_doc.tables[0].cell(0, 1).text, "R1C2")
                self.assertEqual(out_doc.paragraphs[-1].text, "Updated closing")
            finally:
                os.unlink(output_path)
        finally:
            os.unlink(source_path)

    def test_legacy_process_document_can_create_and_process_docx_task_without_task_id(self):
        self._login("writer@conwiz.in")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name

        try:
            doc = Document()
            doc.add_paragraph("Intro paragraph")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A1"
            table.cell(0, 1).text = "B1"
            doc.add_paragraph("Closing paragraph")
            doc.save(source_path)

            with open(source_path, "rb") as infile:
                source_docx_base64 = base64.b64encode(infile.read()).decode("ascii")

            status, payload = self.client.request(
                "POST",
                "/api/process-document",
                {
                    "task_id": "",
                    "source_type": "docx",
                    "source_docx_base64": source_docx_base64,
                    "source_text": "Intro paragraph\nA1\tB1\nClosing paragraph",
                    "source_file_name": "sample.docx",
                    "options": {
                        "spelling": True,
                        "sentence_case": True,
                        "punctuation": True,
                        "chicago_style": True,
                        "ai": {"enabled": False},
                    },
                },
            )
            self.assertEqual(status, 200)
            self.assertTrue(payload.get("success"))
            task_id = payload.get("task_id")
            self.assertTrue(task_id)

            status, task_payload = self.client.request("GET", f"/api/tasks/{task_id}")
            self.assertEqual(status, 200)
            self.assertTrue(task_payload.get("success"))
            task = task_payload.get("task") or {}
            self.assertEqual(task.get("status"), "PROCESSED")
            self.assertTrue(str(task.get("corrected_text") or "").strip())
            self.assertTrue(task.get("downloads", {}).get("clean"))
            self.assertTrue(task.get("downloads", {}).get("highlighted"))
        finally:
            os.unlink(source_path)

    def test_upload_succeeds_when_bridge_header_is_present(self):
        self._login("bridge@conwiz.in")
        status, payload = self.client.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "sample.txt", "content": "Bridge header upload should work."},
            headers={"X-Manuscript-Session": "browser-fallback-session-id"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertTrue(payload.get("task_id"))

    def test_task_access_isolated_between_users(self):
        owner = WsgiTestClient(webapp.app)
        other = WsgiTestClient(webapp.app)

        status, _ = owner.request("POST", "/api/auth/google-login", {"id_token": "test:owner@conwiz.in"})
        self.assertEqual(status, 200)

        status, payload = owner.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "alpha.txt", "content": "Alpha manuscript."},
        )
        self.assertEqual(status, 200)
        task_id = payload.get("task_id")

        status, _ = other.request("POST", "/api/auth/google-login", {"id_token": "test:other@conwiz.in"})
        self.assertEqual(status, 200)

        status, payload = other.request("GET", f"/api/tasks/{task_id}")
        self.assertEqual(status, 404)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "TASK_NOT_FOUND")

    def test_assistant_requires_auth(self):
        status, payload = self.client.request(
            "POST",
            "/api/assistant",
            {"mode": "qna", "message": "help"},
        )
        self.assertEqual(status, 401)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "AUTH_REQUIRED")

    def test_assistant_qna_returns_read_only_diagnostics(self):
        self._login("writer@conwiz.in")
        status, payload = self.client.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "sample.txt", "content": "This are sample text."},
        )
        self.assertEqual(status, 200)
        task_id = str(payload.get("task_id") or "")
        self.assertTrue(task_id)

        status, payload = self.client.request(
            "POST",
            "/api/assistant",
            {"mode": "qna", "task_id": task_id, "message": "show spelling and citation diagnostics"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload.get("mode"), "qna")
        assistant = payload.get("assistant") or {}
        self.assertIsInstance(assistant.get("message"), str)
        self.assertTrue(assistant.get("message"))
        self.assertIsInstance(assistant.get("suggestions"), list)
        diagnostics = assistant.get("task_diagnostics") or {}
        self.assertEqual((diagnostics.get("task") or {}).get("id"), task_id)

    def test_assistant_qna_admin_activity_summary_requires_admin_role(self):
        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request(
            "POST",
            "/api/auth/google-login",
            {"id_token": "test:amit@conwiz.in"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = admin_client.request(
            "POST",
            "/api/assistant",
            {
                "mode": "qna",
                "message": "show all activity",
                "include_admin_activity": True,
            },
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        assistant = payload.get("assistant") or {}
        admin_activity = assistant.get("admin_activity") or {}
        self.assertIsInstance(admin_activity.get("user_counts"), dict)
        self.assertIn("top_event_types", admin_activity)

        member_client = WsgiTestClient(webapp.app)
        status, payload = member_client.request(
            "POST",
            "/api/auth/google-login",
            {"id_token": "test:member@conwiz.in"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        status, payload = member_client.request(
            "POST",
            "/api/assistant",
            {
                "mode": "qna",
                "message": "show all activity",
                "include_admin_activity": True,
            },
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        assistant = payload.get("assistant") or {}
        self.assertNotIn("admin_activity", assistant)

    def test_assistant_action_reprocess_task(self):
        self._login("writer@conwiz.in")
        status, payload = self.client.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "sample.txt", "content": "This are sample text."},
        )
        self.assertEqual(status, 200)
        task_id = str(payload.get("task_id") or "")
        self.assertTrue(task_id)

        status, payload = self.client.request(
            "POST",
            "/api/assistant",
            {
                "mode": "action",
                "action": "reprocess_task",
                "task_id": task_id,
                "options": {"ai": {"enabled": False}},
                "confirm": True,
            },
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload.get("mode"), "action")
        self.assertEqual(payload.get("action"), "reprocess_task")
        result = payload.get("result") or {}
        self.assertTrue(result.get("success"))
        self.assertEqual(result.get("task_id"), task_id)
        self.assertTrue(str(result.get("text") or "").strip())

    def test_assistant_action_reprocess_respects_task_ownership(self):
        owner = WsgiTestClient(webapp.app)
        other = WsgiTestClient(webapp.app)
        status, _ = owner.request("POST", "/api/auth/google-login", {"id_token": "test:owner@conwiz.in"})
        self.assertEqual(status, 200)
        status, payload = owner.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "alpha.txt", "content": "Alpha manuscript."},
        )
        self.assertEqual(status, 200)
        task_id = str(payload.get("task_id") or "")
        self.assertTrue(task_id)

        status, _ = other.request("POST", "/api/auth/google-login", {"id_token": "test:other@conwiz.in"})
        self.assertEqual(status, 200)
        status, payload = other.request(
            "POST",
            "/api/assistant",
            {
                "mode": "action",
                "action": "reprocess_task",
                "task_id": task_id,
                "options": {"ai": {"enabled": False}},
                "confirm": True,
            },
        )
        self.assertEqual(status, 404)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "TASK_NOT_FOUND")

    def test_assistant_action_requires_explicit_confirmation(self):
        self._login("writer@conwiz.in")
        status, payload = self.client.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "sample.txt", "content": "This are sample text."},
        )
        self.assertEqual(status, 200)
        task_id = str(payload.get("task_id") or "")
        self.assertTrue(task_id)

        status, payload = self.client.request(
            "POST",
            "/api/assistant",
            {
                "mode": "action",
                "action": "reprocess_task",
                "task_id": task_id,
                "options": {"ai": {"enabled": False}},
            },
        )
        self.assertEqual(status, 400)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "ASSISTANT_CONFIRMATION_REQUIRED")

    def test_assistant_action_apply_group_decisions(self):
        self._login("writer@conwiz.in")
        status, payload = self.client.request(
            "POST",
            "/api/tasks/upload-text",
            {"file_name": "sample.txt", "content": "This are sample text."},
        )
        self.assertEqual(status, 200)
        task_id = str(payload.get("task_id") or "")
        self.assertTrue(task_id)

        status, payload = self.client.request(
            "POST",
            f"/api/tasks/{task_id}/process",
            {"options": {"ai": {"enabled": False}}},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = self.client.request(
            "POST",
            "/api/assistant",
            {
                "mode": "action",
                "action": "apply_correction_group_decisions",
                "task_id": task_id,
                "group_decisions": {
                    "spelling": True,
                    "capitalization": True,
                    "punctuation": True,
                    "citation": True,
                    "reference": True,
                    "style": True,
                },
                "confirm": True,
            },
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload.get("action"), "apply_correction_group_decisions")
        result = payload.get("result") or {}
        self.assertTrue(result.get("success"))
        self.assertEqual(result.get("task_id"), task_id)

    def test_admin_can_view_users_and_non_admin_cannot(self):
        # Non-admin user should be blocked from admin endpoints.
        self._login("member@conwiz.in")
        status, payload = self.client.request("GET", "/api/admin/users")
        self.assertEqual(status, 403)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "FORBIDDEN")

        # Admin login should allow access.
        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request(
            "POST",
            "/api/auth/google-login",
            {"id_token": "test:amit@conwiz.in"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = admin_client.request("GET", "/api/admin/users")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertGreaterEqual(len(payload.get("users", [])), 1)

    def test_admin_can_validate_ai_provider(self):
        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request(
            "POST",
            "/api/auth/google-login",
            {"id_token": "test:amit@conwiz.in"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = admin_client.request(
            "POST",
            "/api/admin/validate-ai-provider",
            {"provider": "unsupported-provider"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertFalse(payload.get("valid"))
        self.assertIn("Unsupported provider", str(payload.get("message", "")))

    def test_admin_global_settings_round_trip(self):
        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request("POST", "/api/auth/google-login", {"id_token": "test:amit@conwiz.in"})
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = admin_client.request("GET", "/api/admin/global-settings")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertIn("settings", payload)

        updated = {
            "editing": {
                "spelling": True,
                "sentence_case": True,
                "punctuation": True,
                "chicago_style": True,
                "cmos_strict_mode": True,
                "online_reference_validation_admin_cap": 220,
                "auto_resolve_unresolved_references": False,
                "domain_profile": "medical",
                "cmos_profile": "strict",
                "custom_terms": ["myocardial infarction", "HbA1c"],
            },
            "ai": {
                "enabled": True,
                "provider": "agent_router",
                "model": "gpt-5",
                "ollama_host": "http://localhost:11434",
                "gemini_api_key": "gem-key",
                "openrouter_api_key": "or-key",
                "agent_router_api_key": "ar-key",
                "section_wise": True,
                "section_threshold_chars": 14000,
                "section_threshold_paragraphs": 100,
                "section_chunk_chars": 6000,
                "section_chunk_lines": 32,
                "global_consistency_max_chars": 19000,
            },
        }
        status, payload = admin_client.request("POST", "/api/admin/global-settings", {"settings": updated})
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertEqual(payload.get("settings", {}).get("editing", {}).get("domain_profile"), "medical")
        self.assertEqual(payload.get("settings", {}).get("editing", {}).get("cmos_profile"), "strict")
        self.assertEqual(int(payload.get("settings", {}).get("editing", {}).get("online_reference_validation_admin_cap", 0)), 220)
        self.assertFalse(payload.get("settings", {}).get("editing", {}).get("auto_resolve_unresolved_references"))

        user_client = WsgiTestClient(webapp.app)
        status, payload = user_client.request("POST", "/api/auth/google-login", {"id_token": "test:user@conwiz.in"})
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        status, payload = user_client.request("GET", "/api/settings/runtime")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        user_settings = payload.get("settings", {})
        self.assertEqual(user_settings.get("editing", {}).get("domain_profile"), "medical")
        self.assertEqual(user_settings.get("editing", {}).get("cmos_profile"), "strict")
        self.assertEqual(int(user_settings.get("editing", {}).get("online_reference_validation_admin_cap", 0)), 220)
        self.assertFalse(user_settings.get("editing", {}).get("auto_resolve_unresolved_references"))
        self.assertEqual(user_settings.get("ai", {}).get("openrouter_api_key", ""), "")
        self.assertEqual(user_settings.get("ai", {}).get("agent_router_api_key", ""), "")

    def test_admin_reference_validation_diagnostics_requires_admin(self):
        self._login("member@conwiz.in")
        status, payload = self.client.request("GET", "/api/admin/reference-validation-diagnostics")
        self.assertEqual(status, 403)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "FORBIDDEN")

    def test_admin_reference_validation_diagnostics_reset_requires_admin(self):
        self._login("member@conwiz.in")
        status, payload = self.client.request("POST", "/api/admin/reference-validation-diagnostics/reset")
        self.assertEqual(status, 403)
        self.assertFalse(payload.get("success"))
        self.assertEqual(payload.get("error_code"), "FORBIDDEN")

    def test_admin_reference_validation_diagnostics_is_safe_and_structured(self):
        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request("POST", "/api/auth/google-login", {"id_token": "test:amit@conwiz.in"})
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        with patch.dict("os.environ", {"SERPER_API_KEY": "serper-secret-test-key"}, clear=False):
            status, payload = admin_client.request("GET", "/api/admin/reference-validation-diagnostics")

        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        diagnostics = payload.get("diagnostics", {})
        self.assertIsInstance(diagnostics, dict)
        self.assertIn("generated_at", diagnostics)
        self.assertIn("global_runtime", diagnostics)
        self.assertIn("serper", diagnostics)
        self.assertIn("cache", diagnostics)
        self.assertIn("lookup_metrics_last_run", diagnostics)
        self.assertTrue(bool(diagnostics.get("serper", {}).get("configured")))
        serialized = json.dumps(payload)
        self.assertNotIn("serper-secret-test-key", serialized)

    def test_admin_reference_validation_diagnostics_reset_clears_cache(self):
        with ChicagoEditor._SHARED_ONLINE_VALIDATION_CACHE_LOCK:
            ChicagoEditor._SHARED_ONLINE_VALIDATION_CACHE["serper_search:{\"q\":\"test\"}"] = {
                "expires_at": 9999999999,
                "value": [{"title": "cached"}],
            }

        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request("POST", "/api/auth/google-login", {"id_token": "test:amit@conwiz.in"})
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = admin_client.request("POST", "/api/admin/reference-validation-diagnostics/reset")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertGreaterEqual(int(payload.get("removed_cache_entries", 0)), 1)
        diagnostics = payload.get("diagnostics", {})
        self.assertEqual(int((diagnostics.get("cache", {}) or {}).get("entries_total", -1)), 0)

    @patch("chicago_editor.requests.get")
    def test_admin_reference_validation_diagnostics_include_last_run_lookup_metrics(self, mock_get):
        crossref_empty = Mock()
        crossref_empty.status_code = 200
        crossref_empty.raise_for_status.return_value = None
        crossref_empty.json.return_value = {"message": {"items": []}}
        openalex_empty = Mock()
        openalex_empty.status_code = 200
        openalex_empty.raise_for_status.return_value = None
        openalex_empty.json.return_value = {"results": []}
        mock_get.side_effect = [crossref_empty, openalex_empty]

        self._login("writer@conwiz.in")
        status, payload = self.client.request(
            "POST",
            "/api/tasks/upload-text",
            {
                "file_name": "refs.txt",
                "content": (
                    "Introduction cites [1].\n"
                    "References\n"
                    "[1] Kaplan S. The restorative benefits of nature toward an integrative framework. "
                    "J Environ Psychol. 1995 ;15(3):169-182.\n"
                ),
            },
        )
        self.assertEqual(status, 200)
        task_id = payload.get("task_id")
        self.assertTrue(task_id)

        status, payload = self.client.request(
            "POST",
            f"/api/tasks/{task_id}/process",
            {"options": {"ai": {"enabled": False}}},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request("POST", "/api/auth/google-login", {"id_token": "test:amit@conwiz.in"})
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = admin_client.request("GET", "/api/admin/reference-validation-diagnostics")
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        diagnostics = payload.get("diagnostics", {})
        self.assertGreater(int((diagnostics.get("lookup_metrics_last_run", {}) or {}).get("crossref_requests", 0)), 0)
        self.assertGreater(int(diagnostics.get("lookup_metrics_last_run_at", 0) or 0), 0)

    def test_deactivated_user_cannot_access_api(self):
        admin_client = WsgiTestClient(webapp.app)
        user_client = WsgiTestClient(webapp.app)

        status, payload = admin_client.request("POST", "/api/auth/google-login", {"id_token": "test:amit@conwiz.in"})
        self.assertEqual(status, 200)

        status, payload = user_client.request("POST", "/api/auth/google-login", {"id_token": "test:reviewer@conwiz.in"})
        self.assertEqual(status, 200)
        user_id = payload["user"]["id"]

        status, payload = admin_client.request(
            "POST",
            f"/api/admin/users/{user_id}/status",
            {"status": "INACTIVE"},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = user_client.request("GET", "/api/tasks")
        self.assertIn(status, (401, 403))
        self.assertFalse(payload.get("success"))
        self.assertIn(payload.get("error_code"), {"AUTH_USER_INACTIVE", "AUTH_SESSION_INVALID"})

    @patch("webapp.requests.post")
    def test_admin_validate_ai_provider_uses_saved_global_key_when_input_blank(self, mock_post):
        admin_client = WsgiTestClient(webapp.app)
        status, payload = admin_client.request("POST", "/api/auth/google-login", {"id_token": "test:amit@conwiz.in"})
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        status, payload = admin_client.request(
            "POST",
            "/api/admin/global-settings",
            {
                "settings": {
                    "editing": {},
                    "ai": {
                        "provider": "agent_router",
                        "model": "gpt-5",
                        "agent_router_api_key": "saved-agent-token",
                    },
                }
            },
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))

        mock_response = Mock()
        mock_response.status_code = 200
        mock_response.json.return_value = {"choices": [{"message": {"content": "OK"}}]}
        mock_post.return_value = mock_response

        status, payload = admin_client.request(
            "POST",
            "/api/admin/validate-ai-provider",
            {"provider": "agent_router", "model": "", "api_key": "", "ollama_host": ""},
        )
        self.assertEqual(status, 200)
        self.assertTrue(payload.get("success"))
        self.assertTrue(payload.get("valid"))

        _, kwargs = mock_post.call_args
        self.assertEqual(kwargs["headers"]["Authorization"], "Bearer saved-agent-token")
        self.assertEqual(kwargs["json"]["model"], "gpt-5")

    @patch("webapp.requests.get")
    def test_validate_ai_provider_runtime_checks_ollama_model_presence(self, mock_get):
        mock_response = Mock()
        mock_response.status_code = 200
        mock_response.content = b'{"models":[{"name":"llama3.1:latest"}]}'
        mock_response.json.return_value = {"models": [{"name": "llama3.1:latest"}]}
        mock_get.return_value = mock_response

        ok, message = webapp._validate_ai_provider_runtime("ollama", "missing-model", "", "http://localhost:11434")
        self.assertFalse(ok)
        self.assertIn("not installed", message)

        ok, message = webapp._validate_ai_provider_runtime("ollama", "llama3.1:latest", "", "http://localhost:11434")
        self.assertTrue(ok)
        self.assertIn("with model llama3.1:latest", message)


if __name__ == "__main__":
    unittest.main()

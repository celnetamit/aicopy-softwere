"""Backend tests for save/export error codes and runtime telemetry."""

import builtins
import os
import tempfile
import unittest
import zipfile

from docx import Document

import main


class ExportSaveTelemetryTests(unittest.TestCase):
    def setUp(self):
        main.reset_runtime_telemetry()
        main.current_file = "sample.docx"
        main.original_text = ""
        main.corrected_text = ""

    def test_export_without_corrected_text_returns_error_code(self):
        response = main.export_file("clean")
        self.assertFalse(response["success"])
        self.assertEqual(response.get("error_code"), "EXPORT_NO_CORRECTED_DOC")

        telemetry = main.get_runtime_telemetry()["telemetry"]
        self.assertEqual(telemetry["export_attempts"], 1)
        self.assertEqual(telemetry["export_failures"], 1)
        self.assertEqual(telemetry["errors_by_code"].get("EXPORT_NO_CORRECTED_DOC"), 1)

    def test_export_success_updates_telemetry(self):
        main.original_text = "Hello world."
        main.corrected_text = "Hello, world."
        response = main.export_file("clean")
        self.assertTrue(response["success"])
        self.assertIn("base64_data", response)
        self.assertEqual(response["file_name"], "clean_sample.docx")

        telemetry = main.get_runtime_telemetry()["telemetry"]
        self.assertEqual(telemetry["export_attempts"], 1)
        self.assertEqual(telemetry["export_successes"], 1)
        self.assertEqual(telemetry["export_failures"], 0)

    def test_highlighted_export_uses_prefixed_filename(self):
        main.original_text = "Hello world."
        main.corrected_text = "Hello, world."
        response = main.export_file("highlighted")
        self.assertTrue(response["success"])
        self.assertEqual(response["file_name"], "highlighted_sample.docx")

    def test_highlighted_docx_uses_green_insert_and_light_red_delete_text(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            output_path = handle.name

        try:
            main.processor.generate_highlighted_docx("old text", "new text", output_path)
            with zipfile.ZipFile(output_path, "r") as package:
                document_xml = package.read("word/document.xml").decode("utf-8", errors="ignore")
                settings_xml = package.read("word/settings.xml").decode("utf-8", errors="ignore")

            self.assertIn("<w:ins", document_xml)
            self.assertIn("<w:del", document_xml)
            self.assertIn('w:val="2FBF71"', document_xml)
            self.assertIn('w:val="FF9AA8"', document_xml)
            self.assertIn("<w:trackRevisions", settings_xml)
        finally:
            os.unlink(output_path)

    def test_highlighted_docx_does_not_mark_shifted_unchanged_paragraphs_red(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            output_path = handle.name

        original = (
            "First paragraph stays.\n"
            "Second paragraph gets revised.\n"
            "Third paragraph stays the same."
        )
        corrected = (
            "First paragraph stays.\n"
            "A newly inserted paragraph appears here.\n"
            "Second paragraph gets revised carefully.\n"
            "Third paragraph stays the same."
        )

        try:
            main.processor.generate_highlighted_docx(original, corrected, output_path)
            doc = Document(output_path)
            target = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "Third paragraph stays the same.")
            self.assertTrue(target.runs)
            with zipfile.ZipFile(output_path, "r") as package:
                document_xml = package.read("word/document.xml").decode("utf-8", errors="ignore")
            self.assertIn("Third paragraph stays the same.", document_xml)
            self.assertNotIn("Third paragraph stays the same.</w:delText>", document_xml)
            self.assertNotIn("Third paragraph stays the same.</w:t></w:r></w:ins>", document_xml)
        finally:
            os.unlink(output_path)

    def test_highlighted_template_docx_does_not_mark_shifted_unchanged_paragraphs_red(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        source_doc = Document()
        source_doc.add_paragraph("First paragraph stays.")
        source_doc.add_paragraph("Second paragraph gets revised.")
        source_doc.add_paragraph("Third paragraph stays the same.")
        source_doc.save(source_path)

        corrected = (
            "First paragraph stays.\n"
            "A newly inserted paragraph appears here.\n"
            "Second paragraph gets revised carefully.\n"
            "Third paragraph stays the same."
        )

        try:
            main.processor.generate_highlighted_docx("", corrected, output_path, source_docx_path=source_path)
            doc = Document(output_path)
            target = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "Third paragraph stays the same.")
            self.assertTrue(target.runs)
            with zipfile.ZipFile(output_path, "r") as package:
                document_xml = package.read("word/document.xml").decode("utf-8", errors="ignore")
            self.assertIn("Third paragraph stays the same.", document_xml)
            self.assertNotIn("Third paragraph stays the same.</w:delText>", document_xml)
            self.assertNotIn("Third paragraph stays the same.</w:t></w:r></w:ins>", document_xml)
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_highlighted_template_docx_keeps_inserted_paragraph_position(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        source_doc = Document()
        source_doc.add_paragraph("P1")
        source_doc.add_paragraph("P2")
        source_doc.add_paragraph("P3")
        source_doc.save(source_path)

        corrected = "P1\nINSERTED MID\nP2\nP3"

        try:
            main.processor.generate_highlighted_docx("P1\nP2\nP3", corrected, output_path, source_docx_path=source_path)
            doc = Document(output_path)
            texts = [paragraph.text for paragraph in doc.paragraphs]
            self.assertEqual(texts[0], "P1")
            self.assertEqual(texts[2], "P2")
            self.assertEqual(texts[3], "P3")
            with zipfile.ZipFile(output_path, "r") as package:
                document_xml = package.read("word/document.xml").decode("utf-8", errors="ignore")
            self.assertIn("INSERTED MID", document_xml)
            self.assertIn("<w:ins", document_xml)
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_highlighted_template_docx_without_special_parts_keeps_standard_package_layout(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as source_handle:
            source_path = source_handle.name
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as output_handle:
            output_path = output_handle.name

        source_doc = Document()
        source_doc.add_paragraph("Alpha")
        source_doc.add_paragraph("Beta")
        source_doc.save(source_path)

        try:
            main.processor.generate_highlighted_docx("Alpha\nBeta", "Alpha edited\nBeta", output_path, source_docx_path=source_path)
            with zipfile.ZipFile(output_path, "r") as package:
                names = package.namelist()
                self.assertIn("[Content_Types].xml", names)
                self.assertEqual(names[0], "[Content_Types].xml")
        finally:
            os.unlink(source_path)
            os.unlink(output_path)

    def test_clean_docx_renders_missing_placeholders_in_gray(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            output_path = handle.name

        try:
            main.processor.generate_clean_docx("Reference [place missing]", output_path)
            doc = Document(output_path)
            runs = [run for paragraph in doc.paragraphs for run in paragraph.runs if run.text.strip()]

            placeholder_run = next(run for run in runs if run.text == "[place missing]")
            self.assertEqual(str(placeholder_run.font.color.rgb), "808080")
        finally:
            os.unlink(output_path)

    def test_clean_docx_italics_foreign_terms(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            output_path = handle.name

        try:
            main.processor.generate_clean_docx("The remedy was ultra vires and mutatis mutandis applied.", output_path)
            doc = Document(output_path)
            runs = [run for paragraph in doc.paragraphs for run in paragraph.runs if run.text.strip()]

            ultra_vires_run = next(run for run in runs if run.text == "ultra vires")
            mutatis_run = next(run for run in runs if run.text == "mutatis mutandis")
            self.assertTrue(ultra_vires_run.italic)
            self.assertTrue(mutatis_run.italic)
        finally:
            os.unlink(output_path)

    def test_clean_docx_keeps_common_scholarly_latin_in_roman(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            output_path = handle.name

        try:
            main.processor.generate_clean_docx("Method used in vitro and in vivo.", output_path)
            doc = Document(output_path)
            runs = [run for paragraph in doc.paragraphs for run in paragraph.runs if run.text.strip()]
            paragraph_text = " ".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
            self.assertIn("in vitro", paragraph_text)
            self.assertIn("in vivo", paragraph_text)
            self.assertFalse(any(bool(run.italic) and "in vitro" in run.text for run in runs))
            self.assertFalse(any(bool(run.italic) and "in vivo" in run.text for run in runs))
        finally:
            os.unlink(output_path)

    def test_clean_docx_does_not_italicize_foreign_terms_inside_url_or_email(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
            output_path = handle.name

        try:
            main.processor.generate_clean_docx(
                "Use mutatis mutandis in text. URL https://example.com/mutatis Email mutatis@example.com",
                output_path,
            )
            doc = Document(output_path)
            runs = [run for paragraph in doc.paragraphs for run in paragraph.runs if run.text.strip()]

            italic_term_runs = [run for run in runs if run.text == "mutatis mutandis" and bool(run.italic)]
            self.assertEqual(len(italic_term_runs), 1)
        finally:
            os.unlink(output_path)

    def test_save_without_corrected_text_returns_error_code(self):
        response = main.save_file("clean")
        self.assertFalse(response["success"])
        self.assertEqual(response.get("error_code"), "SAVE_NO_CORRECTED_DOC")

        telemetry = main.get_runtime_telemetry()["telemetry"]
        self.assertEqual(telemetry["save_attempts"], 1)
        self.assertEqual(telemetry["save_failures"], 1)
        self.assertEqual(telemetry["errors_by_code"].get("SAVE_NO_CORRECTED_DOC"), 1)

    def test_save_uses_fallback_when_tkinter_missing(self):
        main.original_text = "Hello world."
        main.corrected_text = "Hello, world."

        orig_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name == "tkinter" or name.startswith("tkinter"):
                raise ModuleNotFoundError("No module named 'tkinter'")
            return orig_import(name, *args, **kwargs)

        builtins.__import__ = fake_import
        try:
            response = main.save_file("clean")
        finally:
            builtins.__import__ = orig_import

        self.assertTrue(response["success"])
        self.assertIn("path", response)
        self.assertTrue(os.path.exists(response["path"]))
        self.assertEqual(response.get("warning_code"), "SAVE_DIALOG_UNAVAILABLE_FALLBACK_USED")

        telemetry = main.get_runtime_telemetry()["telemetry"]
        self.assertEqual(telemetry["save_attempts"], 1)
        self.assertEqual(telemetry["save_successes"], 1)
        self.assertEqual(telemetry["save_fallback_used"], 1)

        os.unlink(response["path"])


if __name__ == "__main__":
    unittest.main()
